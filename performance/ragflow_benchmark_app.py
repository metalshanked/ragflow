from __future__ import annotations

import json
import math
import os
import re
import shutil
import statistics
import threading
import time
import uuid
from io import BytesIO
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib
import requests
import urllib3
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from quart import Quart, Response, jsonify, request
from urllib3.exceptions import InsecureRequestWarning

matplotlib.use("Agg")
import matplotlib.pyplot as plt

urllib3.disable_warnings(InsecureRequestWarning)


APP_DIR = Path(__file__).resolve().parent
DATA_DIR = APP_DIR / "data"
RUNS_DIR = DATA_DIR / "runs"
UPLOADS_DIR = DATA_DIR / "uploads"

for directory in (DATA_DIR, RUNS_DIR, UPLOADS_DIR):
    directory.mkdir(parents=True, exist_ok=True)


app = Quart(__name__)
APP_BASE_PATH = ""

RUNS_LOCK = threading.Lock()
RUNS: dict[str, dict[str, Any]] = {}
RUN_SECRETS: dict[str, dict[str, Any]] = {}
BATCH_LIMITERS_LOCK = threading.Lock()
BATCH_LIMITERS: dict[str, threading.Semaphore] = {}
BATCH_ANALYSIS_STATE: dict[str, str] = {}

TERMINAL_DOCUMENT_STATES = {"DONE", "FAIL", "CANCEL"}
TERMINAL_RUN_STATES = {"completed", "failed"}
STOPWORDS = {
    "about", "after", "again", "against", "also", "among", "because", "before",
    "being", "below", "between", "could", "document", "documents", "each", "from",
    "have", "into", "more", "most", "other", "same", "some", "such", "than",
    "that", "their", "them", "then", "there", "these", "they", "this", "those",
    "through", "under", "very", "what", "when", "where", "which", "while", "with",
    "would", "your",
}


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def slugify(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "-", value.strip())
    cleaned = cleaned.strip("-._")
    return cleaned or "file"


def sanitize_filename(name: str) -> str:
    return slugify(Path(name).name)


def normalize_root_url(raw_url: str) -> str:
    url = raw_url.strip().rstrip("/")
    if not url:
        raise ValueError("RAGFlow base URL is required.")
    if url.endswith("/api/v1"):
        return url[: -len("/api/v1")]
    if url.endswith("/v1"):
        return url[: -len("/v1")]
    return url


def normalize_base_path(raw_path: str | None) -> str:
    text = (raw_path or "").strip()
    if not text or text == "/":
        return ""
    return "/" + text.strip("/")


def route_path(path: str) -> str:
    if path in {"", "/"}:
        return APP_BASE_PATH or "/"
    suffix = path if path.startswith("/") else f"/{path}"
    return f"{APP_BASE_PATH}{suffix}" if APP_BASE_PATH else suffix


APP_BASE_PATH = normalize_base_path(os.getenv("PERFORMANCE_APP_BASE_PATH", ""))


def auth_header(api_key: str) -> str:
    api_key = api_key.strip()
    if not api_key:
        raise ValueError("API key is required.")
    return api_key if api_key.lower().startswith("bearer ") else f"Bearer {api_key}"


def safe_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return default
        return float(value)
    except (TypeError, ValueError):
        return default


def safe_int(value: Any, default: int = 0) -> int:
    try:
        if value is None or value == "":
            return default
        return int(value)
    except (TypeError, ValueError):
        return default


def merge_dicts(base: dict[str, Any], updates: dict[str, Any]) -> dict[str, Any]:
    merged = deepcopy(base)
    for key, value in (updates or {}).items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = merge_dicts(merged[key], value)
        else:
            merged[key] = value
    return merged


def parse_json_object(raw_value: str | None, label: str, default: dict[str, Any] | None = None) -> dict[str, Any]:
    text = (raw_value or "").strip()
    if not text:
        return deepcopy(default or {})
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"{label} must be valid JSON: {exc.msg}") from exc
    if not isinstance(parsed, dict):
        raise ValueError(f"{label} must be a JSON object.")
    return parsed


def normalize_dataset_options_for_kb_ui(options: dict[str, Any] | None) -> dict[str, Any]:
    payload = deepcopy(options or {})
    if "chunk_method" in payload and "parser_id" not in payload:
        payload["parser_id"] = payload.pop("chunk_method")
    else:
        payload.pop("chunk_method", None)
    if "embedding_model" in payload and "embd_id" not in payload:
        payload["embd_id"] = payload.pop("embedding_model")
    else:
        payload.pop("embedding_model", None)
    return payload


def normalize_chat_model_name(model_name: Any) -> str | None:
    if model_name is None:
        return None
    text = str(model_name).strip()
    if not text:
        return None
    if text.endswith("@OpenAI-API-Compatible"):
        name, _, provider = text.partition("@")
        if name and not name.endswith("___OpenAI-API"):
            return f"{name}___OpenAI-API@{provider}"
    if text.endswith("@LocalAI"):
        name, _, provider = text.partition("@")
        if name and not name.endswith("___LocalAI"):
            return f"{name}___LocalAI@{provider}"
    if text.endswith("@HuggingFace"):
        name, _, provider = text.partition("@")
        if name and not name.endswith("___HuggingFace"):
            return f"{name}___HuggingFace@{provider}"
    if text.endswith("@VLLM"):
        name, _, provider = text.partition("@")
        if name and not name.endswith("___VLLM"):
            return f"{name}___VLLM@{provider}"
    return text


def normalize_chat_create_options(options: dict[str, Any] | None) -> dict[str, Any]:
    payload = deepcopy(options or {})
    llm = payload.get("llm")
    if isinstance(llm, dict):
        model_name = normalize_chat_model_name(llm.get("model_name"))
        if model_name is None:
            llm.pop("model_name", None)
        else:
            llm["model_name"] = model_name
        if not llm:
            payload.pop("llm", None)
    return payload


def strip_internal_model_suffix(model_name: str) -> str:
    value = model_name or ""
    for suffix in ("___OpenAI-API", "___LocalAI", "___HuggingFace", "___VLLM"):
        if value.endswith(suffix):
            return value[: -len(suffix)]
    return value


def model_name_aliases(model_name: Any) -> set[str]:
    text = str(model_name or "").strip()
    if not text:
        return set()

    aliases = {text}
    base, has_provider, provider = text.partition("@")
    stripped = strip_internal_model_suffix(base)
    aliases.add(stripped)
    if has_provider:
        aliases.add(f"{stripped}@{provider}")
        aliases.add(f"{base}@{provider}")
    return {alias.lower() for alias in aliases if alias}


def model_names_equivalent(left: Any, right: Any) -> bool:
    if left is None or right is None:
        return False
    return bool(model_name_aliases(left) & model_name_aliases(right))


def percentile(values: list[float], pct: float) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    if len(ordered) == 1:
        return ordered[0]
    rank = (len(ordered) - 1) * pct
    lower = math.floor(rank)
    upper = math.ceil(rank)
    if lower == upper:
        return ordered[int(rank)]
    fraction = rank - lower
    return ordered[lower] + (ordered[upper] - ordered[lower]) * fraction


def histogram(values: list[float], bins: int = 10) -> list[dict[str, Any]]:
    if not values:
        return []
    low = min(values)
    high = max(values)
    if math.isclose(low, high):
        return [{"label": f"{low:.1f}", "count": len(values)}]
    step = (high - low) / bins
    counts = [0 for _ in range(bins)]
    for value in values:
        index = min(int((value - low) / step), bins - 1)
        counts[index] += 1
    result = []
    for index, count in enumerate(counts):
        start = low + step * index
        end = start + step
        result.append({"label": f"{start:.0f}-{end:.0f}", "count": count})
    return result


PIPELINE_TS_PATTERNS = [
    re.compile(r"^\s*\[?(\d{4}-\d{2}-\d{2}[ T]\d{2}:\d{2}:\d{2}(?:[.,]\d{1,6})?(?:Z|[+-]\d{2}:?\d{2})?)\]?\s*[:-]?\s*(.*)$"),
    re.compile(r"^\s*\[?(\d{2}:\d{2}:\d{2}(?:[.,]\d{1,6})?)\]?\s*[:-]?\s*(.*)$"),
]


def _coerce_pipeline_label(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        for key in ("label", "name", "title", "component_name", "id"):
            text = str(value.get(key) or "").strip()
            if text:
                return text
    if value is None:
        return ""
    return str(value).strip()


def extract_pipeline_path(dsl: Any) -> list[str]:
    if not isinstance(dsl, dict):
        return []

    path = dsl.get("path")
    graph = dsl.get("graph") if isinstance(dsl.get("graph"), dict) else {}
    nodes = graph.get("nodes") if isinstance(graph.get("nodes"), list) else []
    components = dsl.get("components")

    node_labels: dict[str, str] = {}
    for node in nodes:
        if not isinstance(node, dict):
            continue
        node_id = str(node.get("id") or "").strip()
        data = node.get("data") if isinstance(node.get("data"), dict) else {}
        label = _coerce_pipeline_label(data) or _coerce_pipeline_label(node)
        if node_id and label:
            node_labels[node_id] = label

    if isinstance(components, dict):
        for key, value in components.items():
            label = _coerce_pipeline_label(value) or str(key).strip()
            if label:
                node_labels[str(key).strip()] = label
    elif isinstance(components, list):
        for value in components:
            label = _coerce_pipeline_label(value)
            if label:
                node_labels[label] = label

    resolved: list[str] = []
    if isinstance(path, list):
        for item in path:
            label = _coerce_pipeline_label(item)
            mapped = node_labels.get(label, "")
            final_label = mapped or label
            if final_label and final_label not in resolved:
                resolved.append(final_label)

    if resolved:
        return resolved

    fallback: list[str] = []
    for label in node_labels.values():
        if label and label not in fallback:
            fallback.append(label)
    return fallback[:8]


def parse_pipeline_step_prefix(line: str) -> tuple[datetime | None, str, str]:
    text = line.strip()
    for pattern in PIPELINE_TS_PATTERNS:
        match = pattern.match(text)
        if not match:
            continue
        raw_ts = (match.group(1) or "").strip()
        message = (match.group(2) or "").strip() or "step"
        normalized = raw_ts.replace(",", ".")
        try:
            if "T" in normalized or "-" in normalized:
                if normalized.endswith("Z"):
                    normalized = normalized[:-1] + "+00:00"
                return datetime.fromisoformat(normalized), raw_ts, message
            return datetime.strptime(normalized.ljust(15, "0"), "%H:%M:%S.%f"), raw_ts, message
        except ValueError:
            return None, raw_ts, message
    return None, "", text


def parse_pipeline_progress(progress_msg: str) -> list[dict[str, Any]]:
    lines = [line.strip() for line in (progress_msg or "").splitlines() if line.strip()]
    if not lines:
        return []

    steps: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for line in lines:
        moment, raw_timestamp, message = parse_pipeline_step_prefix(line)
        key = (raw_timestamp or (moment.isoformat() if moment else ""), message)
        if key in seen:
            continue
        seen.add(key)
        steps.append({"message": message, "ts": moment, "timestamp": raw_timestamp, "duration_sec": None, "is_slow": False})

    durations: list[float] = []
    for index in range(len(steps) - 1):
        current = steps[index].get("ts")
        nxt = steps[index + 1].get("ts")
        if isinstance(current, datetime) and isinstance(nxt, datetime):
            elapsed = round(max(0.0, (nxt - current).total_seconds()), 3)
            steps[index]["duration_sec"] = elapsed
            if elapsed > 0:
                durations.append(elapsed)

    threshold = max(percentile(durations, 0.75), 3.0) if durations else 0.0
    if threshold > 0:
        for step in steps:
            duration = safe_float(step.get("duration_sec"))
            step["is_slow"] = duration >= threshold and duration > 0

    normalized: list[dict[str, Any]] = []
    for step in steps:
        normalized.append(
            {
                "message": step["message"],
                "timestamp": step.get("timestamp") or (step["ts"].strftime("%Y-%m-%d %H:%M:%S") if isinstance(step.get("ts"), datetime) else ""),
                "duration_sec": step["duration_sec"],
                "is_slow": bool(step["is_slow"]),
            }
        )
    return normalized


def json_clone(value: Any) -> Any:
    return json.loads(json.dumps(value, ensure_ascii=False))


def chunks_content(chunks: list[dict[str, Any]]) -> str:
    parts = []
    for chunk in chunks:
        content = chunk.get("content") or chunk.get("content_with_weight") or chunk.get("text") or ""
        if content:
            parts.append(content)
    return "\n".join(parts)


def pick_first_sentence(text: str, max_words: int = 14) -> str:
    if not text:
        return ""
    sentences = re.split(r"(?<=[.!?])\s+", re.sub(r"\s+", " ", text).strip())
    for sentence in sentences:
        words = sentence.split()
        if len(words) >= 5:
            return " ".join(words[:max_words]).strip(" ,.;:")
    words = text.split()
    return " ".join(words[:max_words]).strip(" ,.;:")


def top_keywords(text: str, limit: int = 5) -> list[str]:
    words = re.findall(r"\b[A-Za-z][A-Za-z0-9_-]{3,}\b", text.lower())
    counter = Counter(word for word in words if word not in STOPWORDS)
    return [word for word, _ in counter.most_common(limit)]


def build_prompt_set(
    documents: list[dict[str, Any]],
    chunk_samples: dict[str, list[dict[str, Any]]],
    prompts_per_document: int,
    shared_prompts: int,
) -> list[dict[str, Any]]:
    prompts: list[dict[str, Any]] = []
    seen: set[str] = set()

    def add_prompt(kind: str, prompt: str, document: dict[str, Any] | None = None) -> None:
        normalized = prompt.strip()
        if not normalized or normalized in seen:
            return
        seen.add(normalized)
        prompts.append(
            {
                "id": str(uuid.uuid4()),
                "kind": kind,
                "prompt": normalized,
                "document_id": document.get("id") if document else None,
                "document_name": document.get("name") if document else None,
            }
        )

    for document in documents:
        sample_chunks = chunk_samples.get(document["id"], [])
        sample_text = chunks_content(sample_chunks)
        keywords = top_keywords(sample_text)
        sentence = pick_first_sentence(sample_text)

        candidates = [
            ("doc-summary", f"Summarize the most important ideas in the uploaded document '{document['name']}'."),
            ("doc-bullets", f"List the key topics, entities, and claims discussed in '{document['name']}'."),
        ]
        if keywords:
            candidates.append(("doc-keyword", f"What does '{document['name']}' say about {keywords[0]}?"))
        if len(keywords) > 1:
            candidates.append(("doc-keywords", f"Explain how '{document['name']}' covers {keywords[0]}, {keywords[1]}, and any related details."))
        if sentence:
            candidates.append(("doc-grounded", f"Using only the uploaded knowledge, explain this topic in detail: {sentence}."))

        for kind, prompt in candidates[: max(1, prompts_per_document)]:
            add_prompt(kind, prompt, document)

    if len(documents) > 1:
        shared_candidates = [
            ("multi-summary", "Provide an executive summary across all uploaded documents and call out the most important differences."),
            ("multi-compare", "Compare the uploaded documents and explain where they reinforce or contradict each other."),
            ("multi-themes", "Identify the recurring themes across the uploaded documents and support them with concrete details."),
        ]
        for kind, prompt in shared_candidates[: max(0, shared_prompts)]:
            add_prompt(kind, prompt)

    return prompts


class RAGFlowError(RuntimeError):
    pass


class RAGFlowClient:
    def __init__(self, base_url: str, api_key: str, timeout: int = 180, tls_config: dict[str, Any] | None = None):
        self.root = normalize_root_url(base_url)
        self.timeout = timeout
        self.session = requests.Session()
        self.session.headers.update({"Authorization": auth_header(api_key)})
        tls = tls_config or {}
        self.session.verify = bool(tls.get("verify", True))

    def _api_url(self, path: str) -> str:
        return f"{self.root}/api/v1{path}"

    def _web_url(self, path: str) -> str:
        return f"{self.root}/v1{path}"

    def _request(self, method: str, url: str, *, timeout: int | None = None, **kwargs: Any) -> dict[str, Any]:
        response = self.session.request(method, url, timeout=timeout or self.timeout, **kwargs)
        try:
            payload = response.json()
        except ValueError as exc:
            raise RAGFlowError(f"Non-JSON response from {url}: {response.status_code}") from exc

        if response.status_code >= 400:
            raise RAGFlowError(payload.get("message") or response.text or f"HTTP {response.status_code}")
        if isinstance(payload, dict) and "code" in payload and payload.get("code") != 0:
            raise RAGFlowError(payload.get("message") or "RAGFlow request failed")
        return payload

    def create_dataset(self, name: str, options: dict[str, Any] | None = None) -> dict[str, Any]:
        normalized = normalize_dataset_options_for_kb_ui(options)
        if not normalized.get("embd_id"):
            tenant_info = self.get_tenant_info()
            embd_id = tenant_info.get("embd_id")
            if not embd_id:
                raise RAGFlowError("Tenant info did not include a default embd_id for KB create.")
            normalized["embd_id"] = embd_id
        payload = {"name": name}
        payload = merge_dicts(payload, normalized)
        response = self._request("POST", self._web_url("/kb/create"), json=payload)
        kb_id = response.get("data", {}).get("kb_id")
        if not kb_id:
            raise RAGFlowError("KB create response did not include kb_id")
        self.update_dataset(kb_id, **normalized)
        return {"id": kb_id, "name": name}

    def delete_dataset(self, dataset_id: str) -> None:
        self._request("POST", self._web_url("/kb/rm"), json={"kb_id": dataset_id})

    def get_tenant_info(self) -> dict[str, Any]:
        response = self._request("GET", self._web_url("/user/tenant_info"))
        data = response.get("data", {})
        if not isinstance(data, dict):
            raise RAGFlowError("Tenant info response was not an object.")
        return data

    def list_chat_models(self) -> list[str]:
        response = self._request("GET", self._web_url("/llm/list"), params={"model_type": "chat"})
        data = response.get("data", {})
        if not isinstance(data, dict):
            return []
        models: list[str] = []
        for factory, entries in data.items():
            if not isinstance(entries, list):
                continue
            for entry in entries:
                if not isinstance(entry, dict):
                    continue
                llm_name = str(entry.get("llm_name") or entry.get("name") or "").strip()
                fid = str(entry.get("fid") or factory or "").strip()
                if llm_name and fid:
                    models.append(f"{llm_name}@{fid}")
        return models

    def get_default_chat_model_name(self) -> str | None:
        info = self.get_tenant_info()
        value = info.get("llm_id") if isinstance(info, dict) else None
        text = str(value or "").strip()
        return text or None

    def resolve_chat_model_name(self, requested: Any) -> str | None:
        raw_model_name = str(requested or "").strip()
        if not raw_model_name:
            return None
        model_name = normalize_chat_model_name(raw_model_name) or raw_model_name
        default_model = self.get_default_chat_model_name()

        if default_model and model_names_equivalent(model_name, default_model):
            return default_model

        available = self.list_chat_models()
        if not available:
            return default_model if default_model and model_names_equivalent(raw_model_name, default_model) else None

        by_exact = {item.lower(): item for item in available}
        if model_name.lower() in by_exact:
            return by_exact[model_name.lower()]
        if raw_model_name.lower() in by_exact:
            return by_exact[raw_model_name.lower()]

        aliases = {model_name}
        aliases.add(raw_model_name)
        base, has_provider, provider = model_name.partition("@")
        if has_provider:
            stripped = strip_internal_model_suffix(base)
            aliases.add(f"{stripped}@{provider}")
            aliases.add(f"{normalize_chat_model_name(stripped + '@' + provider) or ''}")
        for alias in aliases:
            if alias and alias.lower() in by_exact:
                return by_exact[alias.lower()]

        for item in available:
            item_base, _, item_provider = item.partition("@")
            if provider and item_provider != provider:
                continue
            if strip_internal_model_suffix(item_base).lower() == strip_internal_model_suffix(base).lower():
                return item

        if default_model and model_names_equivalent(raw_model_name, default_model):
            return default_model
        return None

    def get_dataset_detail(self, dataset_id: str) -> dict[str, Any]:
        response = self._request("GET", self._web_url("/kb/detail"), params={"kb_id": dataset_id})
        data = response.get("data", {})
        if not isinstance(data, dict) or not data:
            raise RAGFlowError(f"Dataset detail not found for id={dataset_id}")
        return data

    def update_dataset(self, dataset_id: str, **kwargs: Any) -> dict[str, Any]:
        detail = self.get_dataset_detail(dataset_id)
        payload: dict[str, Any] = {
            "kb_id": dataset_id,
            "name": detail.get("name", ""),
            "description": detail.get("description") or "",
            "parser_id": detail.get("parser_id") or "naive",
        }
        for key in ("embd_id", "permission", "language", "avatar", "pagerank", "pipeline_id"):
            if detail.get(key) is not None:
                payload[key] = detail.get(key)

        updates = deepcopy(kwargs)
        parser_updates = updates.get("parser_config")
        if isinstance(parser_updates, dict) and isinstance(detail.get("parser_config"), dict):
            updates["parser_config"] = merge_dicts(detail["parser_config"], parser_updates)

        payload = merge_dicts(payload, updates)
        payload["kb_id"] = dataset_id
        payload["name"] = str(payload.get("name") or detail.get("name") or "").strip()
        if not payload["name"]:
            raise RAGFlowError(f"Cannot update dataset {dataset_id}: dataset name is empty")
        if payload.get("description") is None:
            payload["description"] = ""
        payload["parser_id"] = payload.get("parser_id") or detail.get("parser_id") or "naive"

        response = self._request("POST", self._web_url("/kb/update"), json=payload)
        data = response.get("data", {})
        return data if isinstance(data, dict) else {}

    def upload_documents(self, dataset_id: str, files: list[Path]) -> list[dict[str, Any]]:
        multipart = []
        handles = []
        try:
            for file_path in files:
                handle = file_path.open("rb")
                handles.append(handle)
                multipart.append(("file", (file_path.name, handle)))
            response = self._request("POST", self._api_url(f"/datasets/{dataset_id}/documents"), files=multipart, timeout=max(self.timeout, 600))
            return response["data"]
        finally:
            for handle in handles:
                handle.close()

    def list_documents(self, dataset_id: str, page: int = 1, page_size: int = 100, document_id: str | None = None) -> dict[str, Any]:
        params = {"page": page, "page_size": page_size}
        if document_id:
            params["id"] = document_id
        response = self._request("GET", self._api_url(f"/datasets/{dataset_id}/documents"), params=params)
        return response["data"]

    def list_all_documents(self, dataset_id: str) -> list[dict[str, Any]]:
        documents = []
        page = 1
        while True:
            payload = self.list_documents(dataset_id, page=page, page_size=100)
            batch = payload.get("docs", [])
            documents.extend(batch)
            if len(batch) < 100:
                break
            page += 1
        return documents

    def parse_documents(self, dataset_id: str, document_ids: list[str], options: dict[str, Any] | None = None) -> None:
        payload = {"document_ids": document_ids}
        payload = merge_dicts(payload, options or {})
        self._request("POST", self._api_url(f"/datasets/{dataset_id}/chunks"), json=payload, timeout=max(self.timeout, 300))

    def list_chunks(self, dataset_id: str, document_id: str, page: int = 1, page_size: int = 10) -> list[dict[str, Any]]:
        response = self._request("GET", self._api_url(f"/datasets/{dataset_id}/documents/{document_id}/chunks"), params={"page": page, "page_size": page_size})
        return response["data"].get("chunks", [])

    def retrieval(self, dataset_id: str, question: str, options: dict[str, Any] | None = None) -> dict[str, Any]:
        config = merge_dicts(
            {
                "top_k": 8,
                "highlight": True,
                "similarity_threshold": 0.2,
                "vector_similarity_weight": 0.7,
            },
            options or {},
        )
        payload = {"dataset_ids": [dataset_id], "question": question}
        payload = merge_dicts(payload, config)
        payload["dataset_ids"] = [dataset_id]
        payload["question"] = question
        payload.setdefault("page", 1)
        payload.setdefault("page_size", payload.get("top_k", 8))
        response = self._request("POST", self._api_url("/retrieval"), json=payload)
        return response["data"]

    def create_chat(self, name: str, dataset_id: str | None, options: dict[str, Any] | None = None) -> dict[str, Any]:
        normalized_options = normalize_chat_create_options(options)
        llm = normalized_options.get("llm")
        llm_id = None
        if isinstance(llm, dict) and llm.get("model_name"):
            requested_model_name = llm.get("model_name")
            default_model_name = self.get_default_chat_model_name()
            if default_model_name and model_names_equivalent(requested_model_name, default_model_name):
                llm.pop("model_name", None)
            else:
                llm_id = self.resolve_chat_model_name(requested_model_name)
                if not llm_id:
                    available_models = self.list_chat_models()
                    available_hint = ", ".join(available_models[:8]) if available_models else "none returned by /v1/llm/list?model_type=chat"
                    raise RAGFlowError(
                        f"Unable to resolve chat model '{requested_model_name}'. "
                        f"Tenant default: {default_model_name or 'unknown'}. "
                        f"Available chat models: {available_hint}"
                    )
            llm.pop("model_name", None)
            if not llm:
                normalized_options.pop("llm", None)
        payload: dict[str, Any] = {"name": name, "dataset_ids": [dataset_id] if dataset_id else []}
        payload = merge_dicts(payload, normalized_options)
        payload["name"] = name
        payload["dataset_ids"] = [dataset_id] if dataset_id else []
        if llm_id:
            payload["llm_id"] = llm_id
        response = self._request("POST", self._api_url("/chats"), json=payload)
        return response["data"]

    def delete_chat(self, chat_id: str) -> None:
        self._request("DELETE", self._api_url("/chats"), json={"ids": [chat_id]})

    def completion(self, chat_id: str, prompt: str, options: dict[str, Any] | None = None) -> dict[str, Any]:
        payload = {
            "model": "ragflow-loadtest",
            "messages": [{"role": "user", "content": prompt}],
            "stream": False,
            "extra_body": {"reference": True},
        }
        payload = merge_dicts(payload, options or {})
        payload["messages"] = [{"role": "user", "content": prompt}]
        return self._request("POST", self._api_url(f"/chats_openai/{chat_id}/chat/completions"), json=payload, timeout=max(self.timeout, 300))

    def list_pipeline_logs(self, dataset_id: str) -> list[dict[str, Any]]:
        response = self._request("POST", self._web_url("/kb/list_pipeline_logs"), params={"kb_id": dataset_id, "page": 1, "page_size": 500}, json={})
        return response["data"].get("logs", [])

    def list_pipeline_dataset_logs(self, dataset_id: str) -> list[dict[str, Any]]:
        response = self._request("POST", self._web_url("/kb/list_pipeline_dataset_logs"), params={"kb_id": dataset_id, "page": 1, "page_size": 500}, json={})
        return response["data"].get("logs", [])


def persist_run_snapshot(run_id: str) -> None:
    with RUNS_LOCK:
        snapshot = json_clone(RUNS[run_id])
    (RUNS_DIR / f"{run_id}.json").write_text(json.dumps(snapshot, indent=2, ensure_ascii=False), encoding="utf-8")


def set_run_fields(run_id: str, **fields: Any) -> None:
    with RUNS_LOCK:
        RUNS[run_id].update(fields)
    persist_run_snapshot(run_id)


def append_run_event(run_id: str, message: str, *, level: str = "info") -> None:
    event = {"ts": utc_now(), "level": level, "message": message}
    with RUNS_LOCK:
        RUNS[run_id]["events"].append(event)
        RUNS[run_id]["events"] = RUNS[run_id]["events"][-200:]
    persist_run_snapshot(run_id)


def patch_run(run_id: str, callback) -> None:
    with RUNS_LOCK:
        callback(RUNS[run_id])
    persist_run_snapshot(run_id)


def set_stage_status(
    run_id: str,
    key: str,
    label: str,
    *,
    status: str,
    started_at: str | None = None,
    completed_at: str | None = None,
    duration_sec: float | None = None,
) -> None:
    def mutate(run: dict[str, Any]) -> None:
        timeline = run.setdefault("timeline", [])
        stage = next((item for item in timeline if item.get("key") == key), None)
        if stage is None:
            stage = {"key": key, "label": label}
            timeline.append(stage)
        stage["label"] = label
        stage["status"] = status
        if started_at is not None:
            stage["started_at"] = started_at
        if completed_at is not None:
            stage["completed_at"] = completed_at
        if duration_sec is not None:
            stage["duration_sec"] = round(max(0.0, safe_float(duration_sec)), 3)

    patch_run(run_id, mutate)


def get_batch_semaphore(batch_id: str, parallel_count: int) -> threading.Semaphore:
    with BATCH_LIMITERS_LOCK:
        semaphore = BATCH_LIMITERS.get(batch_id)
        if semaphore is None:
            semaphore = threading.Semaphore(max(1, parallel_count))
            BATCH_LIMITERS[batch_id] = semaphore
        return semaphore


def list_batch_runs(batch_id: str) -> list[dict[str, Any]]:
    with RUNS_LOCK:
        runs = [json_clone(run) for run in RUNS.values() if run.get("config", {}).get("batch", {}).get("batch_id") == batch_id]
    return sorted(runs, key=lambda item: safe_int(item.get("config", {}).get("batch", {}).get("run_index"), 0))


def set_batch_fields(batch_id: str, **fields: Any) -> None:
    snapshots: list[str] = []
    with RUNS_LOCK:
        for run_id, run in RUNS.items():
            if run.get("config", {}).get("batch", {}).get("batch_id") != batch_id:
                continue
            run.update(fields)
            snapshots.append(run_id)
    for run_id in snapshots:
        persist_run_snapshot(run_id)


def summarize_documents(documents: list[dict[str, Any]], pipeline_logs: list[dict[str, Any]]) -> dict[str, Any]:
    durations = [safe_float(doc.get("process_duration")) for doc in documents if safe_float(doc.get("process_duration")) > 0]
    chunks = [safe_int(doc.get("chunk_count")) for doc in documents]
    tokens = [safe_int(doc.get("token_count")) for doc in documents]
    status_breakdown = Counter(doc.get("run", "UNKNOWN") for doc in documents)
    documents_by_id = {str(doc.get("id") or ""): doc for doc in documents if doc.get("id")}
    pipeline_by_document = []
    pipeline_step_durations: list[float] = []
    pipeline_log_durations: list[float] = []
    seen_document_ids: set[str] = set()

    def build_progress_steps(progress_msg: str, duration_sec: float) -> list[dict[str, Any]]:
        steps = parse_pipeline_progress(progress_msg or "")
        if any(safe_float(step.get("duration_sec")) > 0 for step in steps):
            return steps
        if steps and duration_sec > 0:
            steps[-1]["duration_sec"] = duration_sec
            return steps
        if progress_msg and duration_sec > 0:
            first_line = next((line.strip() for line in progress_msg.splitlines() if line.strip()), "pipeline task")
            return [{"message": first_line, "timestamp": "", "duration_sec": duration_sec, "is_slow": False}]
        return steps

    for log in pipeline_logs:
        document_id = str(log.get("document_id") or "")
        matched_document = documents_by_id.get(document_id, {})
        dsl = log.get("dsl") or {}
        duration_sec = round(max(safe_float(log.get("process_duration")), safe_float(matched_document.get("process_duration"))), 3)
        progress_msg = log.get("progress_msg") or matched_document.get("progress_msg") or ""
        progress_steps = build_progress_steps(progress_msg, duration_sec)
        for step in progress_steps:
            duration = safe_float(step.get("duration_sec"))
            if duration > 0:
                pipeline_step_durations.append(duration)
        if duration_sec > 0:
            pipeline_log_durations.append(duration_sec)
        pipeline_path = extract_pipeline_path(dsl)
        if not pipeline_path:
            fallback_parts = [
                str(log.get("pipeline_title") or "").strip(),
                str(log.get("parser_id") or "").strip(),
                str(matched_document.get("parser_id") or "").strip(),
            ]
            pipeline_path = [part for part in fallback_parts if part]
        pipeline_by_document.append(
            {
                "document_id": document_id or matched_document.get("id"),
                "document_name": log.get("document_name") or matched_document.get("name"),
                "duration_sec": duration_sec,
                "status": log.get("operation_status") or matched_document.get("run"),
                "progress": safe_float(log.get("progress")),
                "progress_msg": progress_msg,
                "task_type": log.get("task_type") or matched_document.get("parser_id") or "",
                "pipeline_path": pipeline_path,
                "progress_steps": progress_steps,
            }
        )
        if document_id:
            seen_document_ids.add(document_id)

    for document in documents:
        document_id = str(document.get("id") or "")
        if not document_id or document_id in seen_document_ids:
            continue
        duration_sec = round(safe_float(document.get("process_duration")), 3)
        progress_msg = document.get("progress_msg") or ""
        if duration_sec <= 0 and not progress_msg:
            continue
        progress_steps = build_progress_steps(progress_msg, duration_sec)
        for step in progress_steps:
            duration = safe_float(step.get("duration_sec"))
            if duration > 0:
                pipeline_step_durations.append(duration)
        if duration_sec > 0:
            pipeline_log_durations.append(duration_sec)
        pipeline_by_document.append(
            {
                "document_id": document_id,
                "document_name": document.get("name"),
                "duration_sec": duration_sec,
                "status": document.get("run"),
                "progress": safe_float(document.get("progress")),
                "progress_msg": progress_msg,
                "task_type": document.get("parser_id") or "parse",
                "pipeline_path": [str(document.get("parser_id") or "parse")],
                "progress_steps": progress_steps,
            }
        )
    total_duration = round(sum(durations), 3)
    total_tokens = sum(tokens)
    return {
        "documents": documents,
        "status_breakdown": dict(status_breakdown),
        "total_documents": len(documents),
        "total_chunks": sum(chunks),
        "total_tokens": total_tokens,
        "total_parse_duration_sec": total_duration,
        "avg_parse_duration_sec": round(statistics.mean(durations), 3) if durations else 0.0,
        "p95_parse_duration_sec": round(percentile(durations, 0.95), 3) if durations else 0.0,
        "token_throughput_per_sec": round(total_tokens / total_duration, 3) if total_duration > 0 else 0.0,
        "pipeline_logs": pipeline_by_document,
        "pipeline_slow_step_threshold_sec": round(max(percentile(pipeline_step_durations, 0.75), 3.0), 3) if pipeline_step_durations else 0.0,
        "pipeline_slow_log_threshold_sec": round(max(percentile(pipeline_log_durations, 0.75), 10.0), 3) if pipeline_log_durations else 0.0,
    }


def summarize_benchmark(results: list[dict[str, Any]], wall_time_sec: float) -> dict[str, Any]:
    latencies = [result["latency_ms"] for result in results if result["ok"]]
    errors = [result for result in results if not result["ok"]]
    prompt_tokens = sum(result.get("prompt_tokens", 0) for result in results if result["ok"])
    completion_tokens = sum(result.get("completion_tokens", 0) for result in results if result["ok"])
    total_tokens = prompt_tokens + completion_tokens
    references = [result.get("reference_count", 0) for result in results if result["ok"]]
    return {
        "count": len(results),
        "ok": len(results) - len(errors),
        "errors": len(errors),
        "error_rate": round(len(errors) / len(results), 4) if results else 0.0,
        "wall_time_sec": round(wall_time_sec, 3),
        "throughput_rps": round(len(results) / wall_time_sec, 3) if wall_time_sec > 0 else 0.0,
        "avg_latency_ms": round(statistics.mean(latencies), 3) if latencies else 0.0,
        "p50_latency_ms": round(percentile(latencies, 0.50), 3) if latencies else 0.0,
        "p95_latency_ms": round(percentile(latencies, 0.95), 3) if latencies else 0.0,
        "p99_latency_ms": round(percentile(latencies, 0.99), 3) if latencies else 0.0,
        "max_latency_ms": round(max(latencies), 3) if latencies else 0.0,
        "min_latency_ms": round(min(latencies), 3) if latencies else 0.0,
        "prompt_tokens": prompt_tokens,
        "completion_tokens": completion_tokens,
        "total_tokens": total_tokens,
        "tokens_per_sec": round(total_tokens / wall_time_sec, 3) if wall_time_sec > 0 else 0.0,
        "avg_reference_count": round(statistics.mean(references), 3) if references else 0.0,
        "latency_histogram": histogram(latencies, bins=10),
    }


def report_safe_text(value: Any) -> str:
    if value is None:
        return "-"
    text = str(value)
    return text if text.strip() else "-"


def chart_label(value: Any, limit: int = 56) -> str:
    text = re.sub(r"\s+", " ", report_safe_text(value)).strip()
    return text if len(text) <= limit else text[: limit - 1] + "..."


def build_horizontal_bar_chart(rows: list[dict[str, Any]], title: str, color: str = "#0f766e") -> BytesIO | None:
    items = [{"label": chart_label(item.get("label")), "value": safe_float(item.get("value"))} for item in rows if safe_float(item.get("value")) > 0]
    if not items:
        return None
    fig_height = max(2.4, 0.45 * len(items) + 1.2)
    fig, ax = plt.subplots(figsize=(8.2, fig_height))
    labels = [item["label"] for item in reversed(items)]
    values = [item["value"] for item in reversed(items)]
    ax.barh(range(len(values)), values, color=color)
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_title(title, fontsize=12)
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def build_histogram_chart(histogram_bins: list[dict[str, Any]], title: str, color: str = "#c2410c") -> BytesIO | None:
    items = [{"label": chart_label(item.get("label"), 20), "count": safe_int(item.get("count"))} for item in histogram_bins if safe_int(item.get("count")) >= 0]
    if not items:
        return None
    fig, ax = plt.subplots(figsize=(8.2, 3.4))
    labels = [item["label"] for item in items]
    counts = [item["count"] for item in items]
    ax.bar(range(len(counts)), counts, color=color)
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=30, ha="right", fontsize=8)
    ax.set_title(title, fontsize=12)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def add_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = report_safe_text(header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = report_safe_text(value)


def add_report_chart(document: Document, image_buffer: BytesIO | None) -> None:
    if image_buffer is None:
        document.add_paragraph("No chart data available.")
        return
    document.add_picture(image_buffer, width=Inches(6.7))


def add_paginated_bar_charts(
    document: Document,
    rows: list[dict[str, Any]],
    title: str,
    *,
    color: str = "#0f766e",
    page_size: int = 12,
) -> None:
    items = [item for item in rows if safe_float(item.get("value")) > 0]
    if not items:
        add_report_chart(document, None)
        return
    for index in range(0, len(items), page_size):
        chunk = items[index : index + page_size]
        chunk_title = title if len(items) <= page_size else f"{title} ({index // page_size + 1}/{math.ceil(len(items) / page_size)})"
        add_report_chart(document, build_horizontal_bar_chart(chunk, chunk_title, color=color))


def build_run_report(run: dict[str, Any]) -> bytes:
    document = Document()
    document.core_properties.title = f"RAGFlow Performance Report - {run.get('id')}"
    document.core_properties.subject = "Performance benchmark results"

    title = document.add_heading("RAGFlow Performance Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Run Overview", level=1)
    add_table(
        document,
        ["Field", "Value"],
        [
            ["Run ID", run.get("id")],
            ["Status", run.get("status")],
            ["Phase", run.get("phase")],
            ["Started", run.get("started_at")],
            ["Completed", run.get("completed_at")],
            ["Dataset", (run.get("dataset") or {}).get("name")],
            ["Dataset ID", (run.get("dataset") or {}).get("id")],
            ["Host", (run.get("config") or {}).get("base_url")],
        ],
    )

    summary = run.get("summary") or {}
    parse = run.get("parse") or {}
    retrieval = run.get("retrieval") or {"summary": {}}
    chat = run.get("chat_results") or {"summary": {}}
    document.add_heading("Key Metrics", level=1)
    add_table(
        document,
        ["Metric", "Value"],
        [
            ["Documents", summary.get("documents", 0)],
            ["Prompts", summary.get("prompts", 0)],
            ["Parse Wall Time (s)", summary.get("parse_wall_time_sec", 0)],
            ["Parse P95 by Document (s)", parse.get("p95_parse_duration_sec", 0)],
            ["Retrieval P95 (ms)", (retrieval.get("summary") or {}).get("p95_latency_ms", 0)],
            ["Chat P95 (ms)", (chat.get("summary") or {}).get("p95_latency_ms", 0)],
            ["Chat Error Rate", f"{safe_float(summary.get('chat_error_rate')) * 100:.2f}%"],
            ["Chat Token Throughput (tok/s)", summary.get("chat_tokens_per_sec", 0)],
        ],
    )

    assessment = ((run.get("analysis") or {}).get("llm_assessment") or {}).get("content")
    if assessment:
        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(assessment)

    parse_rows = [{"label": doc.get("name"), "value": safe_float(doc.get("process_duration"))} for doc in (parse.get("documents") or [])]
    stage_rows = [{"label": stage.get("label") or stage.get("key"), "value": safe_float(stage.get("duration_sec"))} for stage in (run.get("timeline") or [])]
    document.add_heading("Visualizations", level=1)
    document.add_paragraph("Parse Duration by Document")
    add_paginated_bar_charts(document, sorted(parse_rows, key=lambda item: item["value"], reverse=True), "Parse Duration by Document")
    document.add_paragraph("Execution Stage Durations")
    add_paginated_bar_charts(document, sorted(stage_rows, key=lambda item: item["value"], reverse=True), "Execution Stage Durations", color="#c2410c")
    document.add_paragraph("Retrieval Latency Histogram")
    add_report_chart(document, build_histogram_chart((retrieval.get("summary") or {}).get("latency_histogram") or [], "Retrieval Latency Histogram"))
    document.add_paragraph("Chat Latency Histogram")
    add_report_chart(document, build_histogram_chart((chat.get("summary") or {}).get("latency_histogram") or [], "Chat Latency Histogram", color="#0f766e"))

    document.add_heading("Execution Stage Timeline", level=1)
    add_table(
        document,
        ["Stage", "Status", "Duration (s)", "Started", "Completed"],
        [
            [
                stage.get("label") or stage.get("key"),
                stage.get("status"),
                safe_float(stage.get("duration_sec")),
                stage.get("started_at"),
                stage.get("completed_at"),
            ]
            for stage in (run.get("timeline") or [])
        ],
    )

    if parse.get("documents"):
        document.add_heading("Parsed Documents", level=1)
        add_table(
            document,
            ["Document", "Status", "Chunks", "Tokens", "Duration (s)"],
            [
                [
                    doc.get("name"),
                    doc.get("run"),
                    safe_int(doc.get("chunk_count")),
                    safe_int(doc.get("token_count")),
                    safe_float(doc.get("process_duration")),
                ]
                for doc in parse.get("documents", [])
            ],
        )

    prompts = run.get("prompts") or []
    if prompts:
        document.add_heading("Generated Prompts", level=1)
        for item in prompts:
            document.add_paragraph(f"[{item.get('kind')}] {item.get('prompt')}")

    retrieval_results = retrieval.get("results") or []
    if retrieval_results:
        document.add_heading("Retrieval Samples", level=1)
        add_table(
            document,
            ["Prompt", "Latency (ms)", "Chunks", "Top Documents", "Status"],
            [
                [
                    item.get("prompt"),
                    safe_float(item.get("latency_ms")),
                    safe_int(item.get("chunk_count")),
                    ", ".join(item.get("top_documents") or []),
                    "ok" if item.get("ok") else f"error: {item.get('error')}",
                ]
                for item in retrieval_results
            ],
        )

    chat_results = chat.get("results") or []
    if chat_results:
        document.add_heading("Chat Samples", level=1)
        add_table(
            document,
            ["Prompt", "Latency (ms)", "Tokens", "Referenced Docs", "Status"],
            [
                [
                    item.get("prompt"),
                    safe_float(item.get("latency_ms")),
                    safe_int(item.get("total_tokens")),
                    ", ".join(item.get("referenced_documents") or []),
                    "ok" if item.get("ok") else f"error: {item.get('error')}",
                ]
                for item in chat_results
            ],
        )

    events = run.get("events") or []
    if events:
        document.add_heading("Event Log", level=1)
        for event in events:
            document.add_paragraph(f"[{report_safe_text(event.get('ts'))}] {report_safe_text(event.get('level')).upper()} {report_safe_text(event.get('message'))}")

    for section in document.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(10)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_batch_report(batch_id: str, runs: list[dict[str, Any]]) -> bytes:
    ordered_runs = sorted(runs, key=lambda run: safe_int(run.get("config", {}).get("batch", {}).get("run_index"), 0))
    batch_summary = next((run.get("batch_summary") for run in ordered_runs if run.get("batch_summary")), {}) or {}
    aggregate = batch_summary.get("aggregate") or summarize_batch_metrics(ordered_runs)
    llm_summary = ((batch_summary.get("llm_assessment") or {}).get("content")) or ""
    representative = ordered_runs[0] if ordered_runs else {}

    document = Document()
    document.core_properties.title = f"RAGFlow Batch Performance Report - {batch_id}"
    document.core_properties.subject = "Batch performance benchmark results"

    title = document.add_heading("RAGFlow Batch Performance Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Batch Overview", level=1)
    add_table(
        document,
        ["Field", "Value"],
        [
            ["Batch ID", batch_id],
            ["Dataset Prefix", representative.get("config", {}).get("dataset_prefix")],
            ["Host", representative.get("config", {}).get("base_url")],
            ["Requested Parallel Runs", representative.get("config", {}).get("batch", {}).get("parallel_count")],
            ["Total Runs", aggregate.get("total_runs", len(ordered_runs))],
            ["Completed Runs", aggregate.get("completed_runs", 0)],
            ["Failed Runs", aggregate.get("failed_runs", 0)],
            ["Average Parse Wall (s)", aggregate.get("avg_parse_wall_time_sec", 0)],
            ["Average Retrieval P95 (ms)", aggregate.get("avg_retrieval_p95_ms", 0)],
            ["Average Chat P95 (ms)", aggregate.get("avg_chat_p95_ms", 0)],
            ["Average Chat Error Rate", f"{safe_float(aggregate.get('avg_chat_error_rate')) * 100:.2f}%"],
            ["Average Chat Token Throughput (tok/s)", aggregate.get("avg_chat_tokens_per_sec", 0)],
        ],
    )

    if llm_summary:
        document.add_heading("Batch Executive Summary", level=1)
        document.add_paragraph(llm_summary)

    document.add_heading("Batch Run Table", level=1)
    add_table(
        document,
        ["Run", "Status", "Dataset", "Parse Wall (s)", "Retrieval P95 (ms)", "Chat P95 (ms)", "Chat Error Rate"],
        [
            [
                safe_int(run.get("config", {}).get("batch", {}).get("run_index"), 0),
                run.get("status"),
                run.get("dataset", {}).get("name"),
                safe_float(run.get("summary", {}).get("parse_wall_time_sec")),
                safe_float(run.get("summary", {}).get("retrieval_p95_ms")),
                safe_float(run.get("summary", {}).get("chat_p95_ms")),
                f"{safe_float(run.get('summary', {}).get('chat_error_rate')) * 100:.2f}%",
            ]
            for run in ordered_runs
        ],
    )

    for index, run in enumerate(ordered_runs, start=1):
        document.add_page_break()
        run_title = f"Run {safe_int(run.get('config', {}).get('batch', {}).get('run_index'), index)}"
        dataset_name = (run.get("dataset") or {}).get("name")
        if dataset_name:
            run_title = f"{run_title} - {dataset_name}"
        document.add_heading(run_title, level=1)

        add_table(
            document,
            ["Field", "Value"],
            [
                ["Run ID", run.get("id")],
                ["Status", run.get("status")],
                ["Phase", run.get("phase")],
                ["Started", run.get("started_at")],
                ["Completed", run.get("completed_at")],
                ["Dataset", dataset_name],
                ["Dataset ID", (run.get("dataset") or {}).get("id")],
            ],
        )

        summary = run.get("summary") or {}
        parse = run.get("parse") or {}
        retrieval = run.get("retrieval") or {"summary": {}}
        chat = run.get("chat_results") or {"summary": {}}
        document.add_heading("Key Metrics", level=2)
        add_table(
            document,
            ["Metric", "Value"],
            [
                ["Documents", summary.get("documents", 0)],
                ["Prompts", summary.get("prompts", 0)],
                ["Parse Wall Time (s)", summary.get("parse_wall_time_sec", 0)],
                ["Parse P95 by Document (s)", parse.get("p95_parse_duration_sec", 0)],
                ["Retrieval P95 (ms)", (retrieval.get("summary") or {}).get("p95_latency_ms", 0)],
                ["Chat P95 (ms)", (chat.get("summary") or {}).get("p95_latency_ms", 0)],
                ["Chat Error Rate", f"{safe_float(summary.get('chat_error_rate')) * 100:.2f}%"],
                ["Chat Token Throughput (tok/s)", summary.get("chat_tokens_per_sec", 0)],
            ],
        )

        run_assessment = ((run.get("analysis") or {}).get("llm_assessment") or {}).get("content")
        if run_assessment:
            document.add_heading("Run Executive Summary", level=2)
            document.add_paragraph(run_assessment)

        timeline = run.get("timeline") or []
        if timeline:
            document.add_heading("Execution Stage Timeline", level=2)
            add_table(
                document,
                ["Stage", "Status", "Duration (s)", "Started", "Completed"],
                [
                    [
                        stage.get("label") or stage.get("key"),
                        stage.get("status"),
                        safe_float(stage.get("duration_sec")),
                        stage.get("started_at"),
                        stage.get("completed_at"),
                    ]
                    for stage in timeline
                ],
            )

        documents = parse.get("documents") or []
        if documents:
            document.add_heading("Parsed Documents", level=2)
            add_table(
                document,
                ["Document", "Status", "Chunks", "Tokens", "Duration (s)"],
                [
                    [
                        doc.get("name"),
                        doc.get("run"),
                        safe_int(doc.get("chunk_count")),
                        safe_int(doc.get("token_count")),
                        safe_float(doc.get("process_duration")),
                    ]
                    for doc in documents
                ],
            )

        prompts = run.get("prompts") or []
        if prompts:
            document.add_heading("Generated Prompts", level=2)
            for item in prompts:
                document.add_paragraph(f"[{item.get('kind')}] {item.get('prompt')}")

        retrieval_results = retrieval.get("results") or []
        if retrieval_results:
            document.add_heading("Retrieval Samples", level=2)
            add_table(
                document,
                ["Prompt", "Latency (ms)", "Chunks", "Top Documents", "Status"],
                [
                    [
                        item.get("prompt"),
                        safe_float(item.get("latency_ms")),
                        safe_int(item.get("chunk_count")),
                        ", ".join(item.get("top_documents") or []),
                        "ok" if item.get("ok") else f"error: {item.get('error')}",
                    ]
                    for item in retrieval_results
                ],
            )

        chat_results = chat.get("results") or []
        if chat_results:
            document.add_heading("Chat Samples", level=2)
            add_table(
                document,
                ["Prompt", "Latency (ms)", "Tokens", "Referenced Docs", "Status"],
                [
                    [
                        item.get("prompt"),
                        safe_float(item.get("latency_ms")),
                        safe_int(item.get("total_tokens")),
                        ", ".join(item.get("referenced_documents") or []),
                        "ok" if item.get("ok") else f"error: {item.get('error')}",
                    ]
                    for item in chat_results
                ],
            )

    for section in document.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(10)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def summarize_batch_metrics(runs: list[dict[str, Any]]) -> dict[str, Any]:
    parse_times = [safe_float(run.get("summary", {}).get("parse_wall_time_sec")) for run in runs if safe_float(run.get("summary", {}).get("parse_wall_time_sec")) > 0]
    retrieval_p95 = [safe_float(run.get("summary", {}).get("retrieval_p95_ms")) for run in runs if safe_float(run.get("summary", {}).get("retrieval_p95_ms")) > 0]
    chat_p95 = [safe_float(run.get("summary", {}).get("chat_p95_ms")) for run in runs if safe_float(run.get("summary", {}).get("chat_p95_ms")) > 0]
    error_rates = [safe_float(run.get("summary", {}).get("chat_error_rate")) for run in runs if run.get("summary")]
    tokens_per_sec = [safe_float(run.get("summary", {}).get("chat_tokens_per_sec")) for run in runs if safe_float(run.get("summary", {}).get("chat_tokens_per_sec")) > 0]
    statuses = Counter((run.get("status") or "unknown").lower() for run in runs)
    return {
        "total_runs": len(runs),
        "completed_runs": statuses.get("completed", 0),
        "failed_runs": statuses.get("failed", 0),
        "status_breakdown": dict(statuses),
        "avg_parse_wall_time_sec": round(statistics.mean(parse_times), 3) if parse_times else 0.0,
        "p95_parse_wall_time_sec": round(percentile(parse_times, 0.95), 3) if parse_times else 0.0,
        "avg_retrieval_p95_ms": round(statistics.mean(retrieval_p95), 3) if retrieval_p95 else 0.0,
        "max_retrieval_p95_ms": round(max(retrieval_p95), 3) if retrieval_p95 else 0.0,
        "avg_chat_p95_ms": round(statistics.mean(chat_p95), 3) if chat_p95 else 0.0,
        "max_chat_p95_ms": round(max(chat_p95), 3) if chat_p95 else 0.0,
        "avg_chat_error_rate": round(statistics.mean(error_rates), 4) if error_rates else 0.0,
        "avg_chat_tokens_per_sec": round(statistics.mean(tokens_per_sec), 3) if tokens_per_sec else 0.0,
        "runs": [
            {
                "run_id": run.get("id"),
                "run_index": safe_int(run.get("config", {}).get("batch", {}).get("run_index"), 0),
                "status": run.get("status"),
                "dataset_name": run.get("dataset", {}).get("name"),
                "parse_wall_time_sec": safe_float(run.get("summary", {}).get("parse_wall_time_sec")),
                "retrieval_p95_ms": safe_float(run.get("summary", {}).get("retrieval_p95_ms")),
                "chat_p95_ms": safe_float(run.get("summary", {}).get("chat_p95_ms")),
                "chat_error_rate": safe_float(run.get("summary", {}).get("chat_error_rate")),
                "chat_tokens_per_sec": safe_float(run.get("summary", {}).get("chat_tokens_per_sec")),
                "error": run.get("error") or "",
            }
            for run in sorted(runs, key=lambda item: safe_int(item.get("config", {}).get("batch", {}).get("run_index"), 0))
        ],
    }


def build_run_assessment_payload(run: dict[str, Any]) -> dict[str, Any]:
    parse = run.get("parse", {})
    retrieval = run.get("retrieval", {}).get("summary", {})
    chat = run.get("chat_results", {}).get("summary", {})
    slow_logs = sorted(parse.get("pipeline_logs", []), key=lambda item: safe_float(item.get("duration_sec")), reverse=True)[:5]
    return {
        "run_id": run.get("id"),
        "status": run.get("status"),
        "dataset": run.get("dataset", {}),
        "summary": run.get("summary", {}),
        "parse": {
            "documents": parse.get("total_documents", 0),
            "chunks": parse.get("total_chunks", 0),
            "tokens": parse.get("total_tokens", 0),
            "wall_time_sec": parse.get("wall_time_sec", 0),
            "p95_duration_sec": parse.get("p95_parse_duration_sec", 0),
            "slow_log_threshold_sec": parse.get("pipeline_slow_log_threshold_sec", 0),
            "slow_logs": [
                {
                    "document_name": item.get("document_name"),
                    "task_type": item.get("task_type"),
                    "duration_sec": item.get("duration_sec"),
                    "pipeline_path": item.get("pipeline_path", []),
                    "slow_steps": [step for step in item.get("progress_steps", []) if step.get("is_slow")][:5],
                }
                for item in slow_logs
            ],
        },
        "retrieval": retrieval,
        "chat": chat,
    }


def build_batch_assessment_payload(runs: list[dict[str, Any]]) -> dict[str, Any]:
    if not runs:
        return {}
    first = runs[0]
    return {
        "batch": first.get("config", {}).get("batch", {}),
        "host": first.get("config", {}).get("base_url"),
        "aggregate": summarize_batch_metrics(runs),
    }


def build_assessment_prompt(scope: str, payload: dict[str, Any]) -> str:
    subject = "run" if scope == "run" else "batch of runs"
    return (
        "You are analyzing RAGFlow benchmark results.\n"
        f"Focus only on the {subject} metrics in the JSON below.\n"
        "Ignore any retrieved knowledge-base content if present.\n"
        "Write in simple language for an engineer or operator.\n"
        "Cover: what happened, whether the performance looks healthy, the biggest risks or bottlenecks, and 3 concrete next actions.\n"
        "If the data is incomplete, say that explicitly instead of guessing.\n\n"
        f"{json.dumps(payload, indent=2, ensure_ascii=False)}"
    )


def generate_llm_assessment(
    client: RAGFlowClient,
    create_options: dict[str, Any],
    completion_options: dict[str, Any],
    prompt: str,
    name_prefix: str,
) -> dict[str, Any]:
    analysis_options = merge_dicts(
        {
            "prompt": {
                "system": "You analyze benchmark metrics and write concise engineering assessments. Do not say the answer is missing from a knowledge base. Use only the provided benchmark payload.",
                "prologue": "Ready to analyze benchmark results.",
                "parameters": [],
                "empty_response": "No benchmark analysis could be generated.",
                "quote": False,
                "tts": False,
                "refine_multiturn": False,
            }
        },
        normalize_chat_create_options(create_options),
    )
    analysis_chat = client.create_chat(f"{name_prefix}-{datetime.now().strftime('%H%M%S')}", None, options=analysis_options)
    chat_id = analysis_chat["id"]
    try:
        payload = merge_dicts({"extra_body": {"reference": False}}, completion_options or {})
        payload["extra_body"] = merge_dicts(payload.get("extra_body") or {}, {"reference": False})
        response = client.completion(chat_id, prompt, payload)
        choice = ((response.get("choices") or [{}])[0]) if isinstance(response, dict) else {}
        message = choice.get("message") or {}
        return {
            "generated_at": utc_now(),
            "chat_id": chat_id,
            "content": message.get("content") or "",
            "model_name": (analysis_chat.get("llm") or {}).get("model_name") or ((analysis_options.get("llm") or {}).get("model_name")),
        }
    finally:
        try:
            client.delete_chat(chat_id)
        except Exception:
            pass


def maybe_generate_batch_assessment(
    client: RAGFlowClient,
    run_id: str,
    create_options: dict[str, Any],
    completion_options: dict[str, Any],
) -> None:
    with RUNS_LOCK:
        run = RUNS.get(run_id)
        if not run:
            return
        batch = run.get("config", {}).get("batch", {})
        batch_id = batch.get("batch_id")
        if not batch_id:
            return
        runs = [json_clone(item) for item in RUNS.values() if item.get("config", {}).get("batch", {}).get("batch_id") == batch_id]
        if not runs or any((item.get("status") or "").lower() not in TERMINAL_RUN_STATES for item in runs):
            return
        if BATCH_ANALYSIS_STATE.get(batch_id) in {"in_progress", "done"}:
            return
        BATCH_ANALYSIS_STATE[batch_id] = "in_progress"

    try:
        payload = build_batch_assessment_payload(runs)
        if not payload:
            return
        assessment = generate_llm_assessment(
            client,
            create_options,
            completion_options,
            build_assessment_prompt("batch", payload),
            "perf-batch-summary",
        )
        summary = {"aggregate": payload.get("aggregate", {}), "llm_assessment": assessment}
        set_batch_fields(batch_id, batch_summary=summary)
    except Exception:
        with RUNS_LOCK:
            BATCH_ANALYSIS_STATE.pop(batch_id, None)
        raise
    else:
        with RUNS_LOCK:
            BATCH_ANALYSIS_STATE[batch_id] = "done"


def benchmark_retrieval(
    client: RAGFlowClient,
    dataset_id: str,
    prompts: list[dict[str, Any]],
    concurrency: int,
    retrieval_options: dict[str, Any],
    run_id: str,
) -> dict[str, Any]:
    results: list[dict[str, Any]] = []
    started = time.perf_counter()
    total = len(prompts)

    def worker(item: dict[str, Any]) -> dict[str, Any]:
        started_at = time.perf_counter()
        try:
            payload = client.retrieval(dataset_id, item["prompt"], retrieval_options)
            elapsed_ms = round((time.perf_counter() - started_at) * 1000, 3)
            chunks = payload.get("chunks", [])
            doc_aggs = payload.get("doc_aggs", [])
            return {
                "prompt_id": item["id"],
                "prompt": item["prompt"],
                "kind": item["kind"],
                "document_name": item.get("document_name"),
                "ok": True,
                "latency_ms": elapsed_ms,
                "chunk_count": len(chunks),
                "top_documents": [agg.get("doc_name") for agg in doc_aggs[:3] if agg.get("doc_name")],
                "top_chunk_preview": (chunks[0].get("content_with_weight") or chunks[0].get("content") or "")[:240] if chunks else "",
                "error": "",
            }
        except Exception as exc:
            elapsed_ms = round((time.perf_counter() - started_at) * 1000, 3)
            return {
                "prompt_id": item["id"],
                "prompt": item["prompt"],
                "kind": item["kind"],
                "document_name": item.get("document_name"),
                "ok": False,
                "latency_ms": elapsed_ms,
                "chunk_count": 0,
                "top_documents": [],
                "top_chunk_preview": "",
                "error": str(exc),
            }

    with ThreadPoolExecutor(max_workers=max(1, concurrency)) as pool:
        future_map = {pool.submit(worker, item): item for item in prompts}
        completed = 0
        for future in as_completed(future_map):
            result = future.result()
            results.append(result)
            completed += 1
            patch_run(
                run_id,
                lambda run: run["live"].update(
                    {
                        "retrieval_completed": completed,
                        "retrieval_total": total,
                        "last_retrieval": {
                            "prompt": result["prompt"][:120],
                            "ok": result["ok"],
                            "latency_ms": result["latency_ms"],
                        },
                    }
                ),
            )

    results.sort(key=lambda item: item["prompt"])
    return {"results": results, "summary": summarize_benchmark(results, time.perf_counter() - started)}


def benchmark_chat(
    client: RAGFlowClient,
    chat_id: str,
    prompts: list[dict[str, Any]],
    concurrency: int,
    completion_options: dict[str, Any],
    run_id: str,
) -> dict[str, Any]:
    results: list[dict[str, Any]] = []
    started = time.perf_counter()
    total = len(prompts)

    def worker(item: dict[str, Any]) -> dict[str, Any]:
        started_at = time.perf_counter()
        try:
            payload = client.completion(chat_id, item["prompt"], options=completion_options)
            elapsed_ms = round((time.perf_counter() - started_at) * 1000, 3)
            choice = (payload.get("choices") or [{}])[0]
            message = choice.get("message") or {}
            usage = payload.get("usage") or {}
            reference = message.get("reference") or []
            answer = message.get("content") or ""
            return {
                "prompt_id": item["id"],
                "prompt": item["prompt"],
                "kind": item["kind"],
                "document_name": item.get("document_name"),
                "ok": True,
                "latency_ms": elapsed_ms,
                "prompt_tokens": safe_int(usage.get("prompt_tokens")),
                "completion_tokens": safe_int(usage.get("completion_tokens")),
                "total_tokens": safe_int(usage.get("total_tokens")),
                "reference_count": len(reference),
                "referenced_documents": list(
                    {
                        chunk.get("document_name") or chunk.get("doc_name") or chunk.get("document_id")
                        for chunk in reference
                        if chunk.get("document_name") or chunk.get("doc_name") or chunk.get("document_id")
                    }
                ),
                "answer_preview": answer[:360],
                "finish_reason": choice.get("finish_reason"),
                "error": "",
            }
        except Exception as exc:
            elapsed_ms = round((time.perf_counter() - started_at) * 1000, 3)
            return {
                "prompt_id": item["id"],
                "prompt": item["prompt"],
                "kind": item["kind"],
                "document_name": item.get("document_name"),
                "ok": False,
                "latency_ms": elapsed_ms,
                "prompt_tokens": 0,
                "completion_tokens": 0,
                "total_tokens": 0,
                "reference_count": 0,
                "referenced_documents": [],
                "answer_preview": "",
                "finish_reason": "error",
                "error": str(exc),
            }

    with ThreadPoolExecutor(max_workers=max(1, concurrency)) as pool:
        future_map = {pool.submit(worker, item): item for item in prompts}
        completed = 0
        for future in as_completed(future_map):
            result = future.result()
            results.append(result)
            completed += 1
            patch_run(
                run_id,
                lambda run: run["live"].update(
                    {
                        "chat_completed": completed,
                        "chat_total": total,
                        "last_chat": {
                            "prompt": result["prompt"][:120],
                            "ok": result["ok"],
                            "latency_ms": result["latency_ms"],
                        },
                    }
                ),
            )

    results.sort(key=lambda item: item["prompt"])
    return {"results": results, "summary": summarize_benchmark(results, time.perf_counter() - started)}


def cleanup_remote_resources(client: RAGFlowClient, resource_ids: dict[str, str], run_id: str) -> None:
    chat_id = resource_ids.get("chat_id")
    dataset_id = resource_ids.get("dataset_id")
    if chat_id:
        try:
            client.delete_chat(chat_id)
            append_run_event(run_id, f"Deleted remote chat assistant {chat_id}.")
        except Exception as exc:
            append_run_event(run_id, f"Chat cleanup failed: {exc}", level="warning")
    if dataset_id:
        try:
            client.delete_dataset(dataset_id)
            append_run_event(run_id, f"Deleted remote dataset {dataset_id}.")
        except Exception as exc:
            append_run_event(run_id, f"Dataset cleanup failed: {exc}", level="warning")


def execute_run(run_id: str) -> None:
    with RUNS_LOCK:
        config = deepcopy(RUN_SECRETS[run_id])

    client = RAGFlowClient(config["base_url"], config["api_key"], timeout=config["http_timeout_sec"], tls_config=config.get("tls"))
    resource_ids: dict[str, str] = {}
    batch_semaphore: threading.Semaphore | None = None
    active_stage: dict[str, Any] | None = None

    def start_stage(key: str, label: str) -> None:
        nonlocal active_stage
        active_stage = {"key": key, "label": label, "started_at": utc_now(), "started_perf": time.perf_counter()}
        set_stage_status(run_id, key, label, status="running", started_at=active_stage["started_at"])

    def finish_stage(status: str = "completed") -> None:
        nonlocal active_stage
        if not active_stage:
            return
        set_stage_status(
            run_id,
            active_stage["key"],
            active_stage["label"],
            status=status,
            completed_at=utc_now(),
            duration_sec=time.perf_counter() - active_stage["started_perf"],
        )
        active_stage = None

    try:
        batch = config.get("batch", {})
        batch_id = batch.get("batch_id")
        parallel_count = max(1, safe_int(batch.get("parallel_count"), 1))
        if batch_id:
            batch_semaphore = get_batch_semaphore(batch_id, parallel_count)
            if safe_int(batch.get("run_count"), 1) > parallel_count:
                set_run_fields(run_id, phase="waiting_for_slot")
                append_run_event(run_id, f"Waiting for a batch execution slot ({parallel_count} parallel run(s) allowed).")
            batch_semaphore.acquire()
        set_run_fields(run_id, status="running", phase="provisioning", started_at=utc_now())
        append_run_event(run_id, f"Run started against {normalize_root_url(config['base_url'])}.")

        run_suffix = f"-r{safe_int(batch.get('run_index'), 1):02d}" if safe_int(batch.get("run_count"), 1) > 1 else ""
        dataset_name = f"{config['dataset_prefix']}-{datetime.now().strftime('%Y%m%d-%H%M%S')}{run_suffix}"
        dataset_options = merge_dicts(
            {
                "chunk_method": "naive",
                "description": "Generated by performance load tester",
                "parser_config": {"layout_recognize": "DeepDOC"},
            },
            config["dataset_options"],
        )
        start_stage("provisioning", "Provisioning")
        dataset = client.create_dataset(dataset_name, dataset_options)
        dataset_id = dataset["id"]
        resource_ids["dataset_id"] = dataset_id
        set_run_fields(run_id, dataset={"id": dataset_id, "name": dataset_name})
        append_run_event(run_id, f"Created dataset {dataset_name} ({dataset_id}).")
        set_run_fields(run_id, progress=10)
        finish_stage()

        start_stage("upload", "Upload Documents")
        uploaded_documents = client.upload_documents(dataset_id, [Path(path) for path in config["local_files"]])
        uploaded_ids = [document["id"] for document in uploaded_documents]
        patch_run(run_id, lambda run: run["documents"].update({"uploaded": uploaded_documents}))
        append_run_event(run_id, f"Uploaded {len(uploaded_documents)} document(s).")
        set_run_fields(run_id, progress=20)
        finish_stage()

        enable_parsing = config["stages"]["parsing"]["enabled"]
        enable_retrieval = config["stages"]["retrieval"]["enabled"]
        enable_chat = config["stages"]["chat"]["enabled"]

        parse_stage = config["stages"]["parsing"]
        retrieval_stage = config["stages"]["retrieval"]
        chat_stage = config["stages"]["chat"]

        parse_summary: dict[str, Any] = {}
        last_snapshot: list[dict[str, Any]] = []
        parse_wall_time = 0.0

        if enable_parsing:
            set_run_fields(run_id, phase="parsing")
            start_stage("parsing", "Parsing")
            client.parse_documents(dataset_id, uploaded_ids, parse_stage["request_options"])
            append_run_event(run_id, "Document parsing queued.")

            parse_started = time.perf_counter()
            while True:
                documents = client.list_all_documents(dataset_id)
                tracked = [document for document in documents if document["id"] in uploaded_ids]
                last_snapshot = tracked
                done_count = sum(1 for document in tracked if document.get("run") in TERMINAL_DOCUMENT_STATES)
                patch_run(run_id, lambda run: run["live"].update({"parse_completed": done_count, "parse_total": len(uploaded_ids), "documents": tracked}))
                if tracked and done_count == len(uploaded_ids):
                    break
                time.sleep(parse_stage["poll_interval_sec"])

            parse_wall_time = round(time.perf_counter() - parse_started, 3)
            file_logs = client.list_pipeline_logs(dataset_id) if parse_stage["collect_pipeline_logs"] else []
            dataset_logs = client.list_pipeline_dataset_logs(dataset_id) if parse_stage["collect_pipeline_logs"] else []
            parse_summary = summarize_documents(last_snapshot, file_logs)
            parse_summary["wall_time_sec"] = parse_wall_time
            parse_summary["dataset_logs"] = dataset_logs
            set_run_fields(run_id, parse=parse_summary, progress=50)
            append_run_event(run_id, "Parsing completed and pipeline logs collected.")
            finish_stage()
        else:
            last_snapshot = client.list_all_documents(dataset_id)
            set_run_fields(run_id, phase="uploaded", parse={"skipped": True, "documents": last_snapshot}, progress=40)
            append_run_event(run_id, "Parsing stage skipped.")
            set_stage_status(run_id, "parsing", "Parsing", status="skipped", duration_sec=0.0)

        prompts: list[dict[str, Any]] = []
        if enable_retrieval or enable_chat:
            set_run_fields(run_id, phase="prompt_generation")
            start_stage("prompt_generation", "Prompt Generation")
            chunk_samples = {}
            sample_size = max(3, parse_stage["chunk_sample_size"], config["prompts_per_document"])
            for document in last_snapshot:
                if document.get("run") == "DONE":
                    chunk_samples[document["id"]] = client.list_chunks(dataset_id, document["id"], page=1, page_size=sample_size)
            prompts = build_prompt_set(last_snapshot, chunk_samples, prompts_per_document=config["prompts_per_document"], shared_prompts=config["shared_prompts"])
            if not prompts:
                prompts = [{"id": str(uuid.uuid4()), "kind": "fallback", "prompt": "Summarize the uploaded knowledge base and highlight the most important facts.", "document_id": None, "document_name": None}]
            set_run_fields(run_id, prompts=prompts, progress=55)
            append_run_event(run_id, f"Generated {len(prompts)} automatic prompt(s).")
            finish_stage()
        else:
            append_run_event(run_id, "Prompt generation skipped because retrieval and chat are disabled.")
            set_stage_status(run_id, "prompt_generation", "Prompt Generation", status="skipped", duration_sec=0.0)

        retrieval: dict[str, Any] = {"summary": {"skipped": not enable_retrieval}, "results": []}
        chat_results: dict[str, Any] = {"summary": {"skipped": not enable_chat}, "results": []}

        if enable_retrieval:
            set_run_fields(run_id, phase="retrieval_benchmark", progress=65)
            start_stage("retrieval", "Retrieval Benchmark")
            retrieval = benchmark_retrieval(
                client,
                dataset_id,
                prompts,
                retrieval_stage["concurrency"],
                retrieval_stage["request_options"],
                run_id,
            )
            set_run_fields(run_id, retrieval=retrieval, progress=78)
            append_run_event(run_id, "Retrieval benchmark completed.")
            finish_stage()
        else:
            set_run_fields(run_id, retrieval=retrieval)
            append_run_event(run_id, "Retrieval stage skipped.")
            set_stage_status(run_id, "retrieval", "Retrieval Benchmark", status="skipped", duration_sec=0.0)

        if enable_chat:
            set_run_fields(run_id, phase="assistant_setup", progress=80)
            start_stage("assistant_setup", "Assistant Setup")
            chat_name = f"{config['chat_prefix']}-{datetime.now().strftime('%H%M%S')}{run_suffix}"
            chat = client.create_chat(chat_name, dataset_id, options=chat_stage["create_options"])
            chat_id = chat["id"]
            resource_ids["chat_id"] = chat_id
            set_run_fields(run_id, chat={"id": chat_id, "name": chat_name})
            append_run_event(run_id, f"Created chat assistant {chat_name} ({chat_id}).")
            finish_stage()

            set_run_fields(run_id, phase="chat_benchmark", progress=85)
            start_stage("chat", "Chat Benchmark")
            chat_results = benchmark_chat(
                client,
                chat_id,
                prompts,
                chat_stage["concurrency"],
                chat_stage["completion_options"],
                run_id,
            )
            set_run_fields(run_id, chat_results=chat_results, progress=95)
            append_run_event(run_id, "Chat benchmark completed.")
            finish_stage()
        else:
            set_run_fields(run_id, chat_results=chat_results)
            append_run_event(run_id, "Chat stage skipped.")
            set_stage_status(run_id, "assistant_setup", "Assistant Setup", status="skipped", duration_sec=0.0)
            set_stage_status(run_id, "chat", "Chat Benchmark", status="skipped", duration_sec=0.0)

        final_summary = {
            "documents": parse_summary.get("total_documents", len(last_snapshot)),
            "prompts": len(prompts),
            "parse_wall_time_sec": parse_wall_time,
            "parse_total_tokens": parse_summary.get("total_tokens", 0),
            "retrieval_p95_ms": retrieval.get("summary", {}).get("p95_latency_ms", 0),
            "chat_p95_ms": chat_results.get("summary", {}).get("p95_latency_ms", 0),
            "chat_error_rate": chat_results.get("summary", {}).get("error_rate", 0),
            "chat_tokens_per_sec": chat_results.get("summary", {}).get("tokens_per_sec", 0),
        }
        set_run_fields(run_id, summary=final_summary)

        if config["analysis"]["enabled"]:
            try:
                set_run_fields(run_id, phase="llm_summary", progress=97)
                start_stage("llm_summary", "Executive Summary")
                run_payload = build_run_response(run_id)
                assessment = generate_llm_assessment(
                    client,
                    chat_stage["create_options"],
                    chat_stage["completion_options"],
                    build_assessment_prompt("run", build_run_assessment_payload(run_payload)),
                    "perf-run-summary",
                )
                set_run_fields(run_id, analysis={"enabled": True, "llm_assessment": assessment})
                append_run_event(run_id, "Generated executive summary with the configured RAGFlow LLM.")
                finish_stage()
            except Exception as exc:
                finish_stage("failed")
                set_run_fields(run_id, analysis={"enabled": True, "error": str(exc)})
                append_run_event(run_id, f"Executive summary skipped: {exc}", level="warning")
        else:
            set_run_fields(run_id, analysis={"enabled": False, "skipped": True})
            set_stage_status(run_id, "llm_summary", "Executive Summary", status="skipped", duration_sec=0.0)

        set_run_fields(
            run_id,
            summary=final_summary,
            progress=100,
            phase="completed",
            status="completed",
            completed_at=utc_now(),
        )
        append_run_event(run_id, "Run completed successfully.")
        if config["analysis"]["enabled"] and resource_ids.get("dataset_id"):
            try:
                maybe_generate_batch_assessment(
                    client,
                    run_id,
                    chat_stage["create_options"],
                    chat_stage["completion_options"],
                )
            except Exception as exc:
                append_run_event(run_id, f"Batch executive summary skipped: {exc}", level="warning")
    except Exception as exc:
        finish_stage("failed")
        set_run_fields(run_id, status="failed", phase="failed", completed_at=utc_now(), error=str(exc))
        append_run_event(run_id, f"Run failed: {exc}", level="error")
    finally:
        if batch_semaphore is not None:
            batch_semaphore.release()
        if config.get("cleanup_remote"):
            start_stage("cleanup", "Cleanup Remote Resources")
            cleanup_remote_resources(client, resource_ids, run_id)
            finish_stage()
        else:
            set_stage_status(run_id, "cleanup", "Cleanup Remote Resources", status="skipped", duration_sec=0.0)
        upload_dir = Path(config["upload_dir"])
        if upload_dir.exists():
            shutil.rmtree(upload_dir, ignore_errors=True)
        batch_root_raw = config.get("batch_root")
        if batch_root_raw:
            batch_root = Path(batch_root_raw)
            try:
                batch_root.rmdir()
            except OSError:
                pass
        with RUNS_LOCK:
            RUN_SECRETS.pop(run_id, None)
        persist_run_snapshot(run_id)


def build_run_response(run_id: str) -> dict[str, Any]:
    with RUNS_LOCK:
        return json_clone(RUNS[run_id])


def queue_run(run_id: str, config: dict[str, Any], filenames: list[str]) -> None:
    batch = config.get("batch", {})
    run_count = safe_int(batch.get("run_count"), 1)
    run_index = safe_int(batch.get("run_index"), 1)
    batch_note = f" ({run_index}/{run_count})" if run_count > 1 else ""
    run = {
        "id": run_id,
        "created_at": utc_now(),
        "started_at": None,
        "completed_at": None,
        "status": "queued",
        "phase": "queued",
        "progress": 0,
        "error": "",
        "config": {
            "base_url": normalize_root_url(config["base_url"]),
            "api_key_hint": config["api_key_hint"],
            "tls": {"verify_ssl": bool(config["tls"].get("verify", True))},
            "dataset_options": config["dataset_options"],
            "prompts_per_document": config["prompts_per_document"],
            "shared_prompts": config["shared_prompts"],
            "dataset_prefix": config["dataset_prefix"],
            "chat_prefix": config["chat_prefix"],
            "cleanup_remote": config["cleanup_remote"],
            "stages": config["stages"],
            "analysis": config["analysis"],
            "batch": batch,
        },
        "dataset": {},
        "chat": {},
        "documents": {"filenames": filenames, "uploaded": []},
        "parse": {},
        "prompts": [],
        "retrieval": {},
        "chat_results": {},
        "analysis": {},
        "batch_summary": {},
        "summary": {},
        "timeline": [],
        "live": {"parse_completed": 0, "parse_total": len(filenames), "retrieval_completed": 0, "retrieval_total": 0, "chat_completed": 0, "chat_total": 0, "documents": []},
        "events": [{"ts": utc_now(), "level": "info", "message": f"Run queued{batch_note} with {len(filenames)} file(s)."}],
    }

    with RUNS_LOCK:
        RUNS[run_id] = run
        RUN_SECRETS[run_id] = config
    persist_run_snapshot(run_id)
    threading.Thread(target=execute_run, args=(run_id,), daemon=True).start()


def load_existing_runs() -> None:
    for json_file in sorted(RUNS_DIR.glob("*.json")):
        try:
            payload = json.loads(json_file.read_text(encoding="utf-8"))
            run_id = payload.get("id")
            if run_id:
                if (payload.get("status") or "").lower() in {"queued", "running"}:
                    payload["status"] = "failed"
                    payload["phase"] = "interrupted"
                    payload["error"] = payload.get("error") or "The performance app restarted before this run finished."
                    payload["completed_at"] = payload.get("completed_at") or utc_now()
                RUNS[run_id] = payload
        except Exception:
            continue

INDEX_HTML = """<!doctype html>
<html lang=\"en\">
<head>
  <meta charset=\"utf-8\" />
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
  <title>RAGFlow Performance Lab</title>
  <link rel=\"icon\" type=\"image/svg+xml\" href=\"data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='12' fill='%230f766e'/%3E%3Crect x='14' y='34' width='8' height='16' rx='2' fill='%23fffaf0'/%3E%3Crect x='28' y='24' width='8' height='26' rx='2' fill='%23fffaf0'/%3E%3Crect x='42' y='16' width='8' height='34' rx='2' fill='%23fffaf0'/%3E%3C/svg%3E\" />
  <style>
    :root { --bg:#f4efe6; --panel:#fffaf0; --ink:#1d232a; --muted:#6d726f; --line:#d9cfbe; --accent:#0f766e; --accent2:#c2410c; --good:#166534; --bad:#b91c1c; }
    * { box-sizing:border-box; }
    body { margin:0; background:radial-gradient(circle at top left, rgba(15,118,110,.12), transparent 28%), radial-gradient(circle at top right, rgba(194,65,12,.12), transparent 24%), var(--bg); color:var(--ink); font-family:\"Segoe UI\",sans-serif; }
    .page { display:grid; grid-template-columns:380px minmax(0,1fr); min-height:100vh; }
    .sidebar, .content { padding:20px; }
    .sidebar { border-right:1px solid var(--line); background:rgba(255,250,240,.82); backdrop-filter:blur(8px); }
    .panel, .brand, .card { border:1px solid var(--line); background:var(--panel); border-radius:18px; box-shadow:0 6px 24px rgba(29,35,42,.05); }
    .brand, .panel, .card { padding:16px; }
    .brand { margin-bottom:16px; }
    .brand h1 { margin:0 0 6px; font-size:26px; letter-spacing:-.04em; }
    .brand p, .subtle { color:var(--muted); line-height:1.45; }
    .panel { margin-bottom:16px; }
    h2, h3 { margin:0 0 12px; letter-spacing:-.02em; }
    .form-grid { display:grid; gap:12px; }
    .row2 { display:grid; gap:12px; grid-template-columns:repeat(2,minmax(0,1fr)); }
    label { display:grid; gap:6px; font-size:13px; color:var(--muted); }
    input, select, textarea, button { font:inherit; }
    input, select, textarea { width:100%; border:1px solid var(--line); background:#fff; border-radius:12px; padding:10px 12px; color:var(--ink); }
    textarea { min-height:96px; resize:vertical; font-family:Consolas,monospace; font-size:12px; }
    button { border:0; border-radius:12px; padding:12px 14px; cursor:pointer; background:var(--accent); color:white; font-weight:600; }
    button.secondary { background:#e7ddd0; color:var(--ink); }
    .check { display:flex; align-items:center; gap:10px; color:var(--ink); }
    .check input { width:auto; }
    .run-list { display:grid; gap:10px; max-height:38vh; overflow:auto; }
    .run-item { border:1px solid var(--line); border-radius:14px; padding:12px; background:#fff; cursor:pointer; }
    .run-item.active { border-color:var(--accent); box-shadow:inset 0 0 0 1px var(--accent); }
    .meta { display:flex; justify-content:space-between; gap:8px; font-size:12px; color:var(--muted); }
    .badge { display:inline-flex; padding:4px 10px; border-radius:999px; font-size:12px; font-weight:700; text-transform:uppercase; letter-spacing:.04em; }
    .badge.completed { background:rgba(22,101,52,.1); color:var(--good); }
    .badge.running { background:rgba(15,118,110,.1); color:var(--accent); }
    .badge.failed { background:rgba(185,28,28,.1); color:var(--bad); }
    .grid { display:grid; gap:16px; grid-template-columns:repeat(12,minmax(0,1fr)); }
    .span3 { grid-column:span 3; } .span6 { grid-column:span 6; } .span12 { grid-column:span 12; }
    .tabs { display:flex; flex-wrap:wrap; gap:8px; margin:0 0 16px; }
    .tab { border:1px solid var(--line); border-radius:999px; padding:8px 12px; background:#fff; color:var(--ink); cursor:pointer; }
    .tab.active { border-color:var(--accent); color:var(--accent); box-shadow:inset 0 0 0 1px var(--accent); }
    .metric .label { color:var(--muted); font-size:12px; text-transform:uppercase; letter-spacing:.06em; }
    .metric .value { font-size:28px; font-weight:800; letter-spacing:-.05em; }
    .progress { height:12px; border-radius:999px; background:#e8dfd4; overflow:hidden; margin:10px 0 8px; }
    .progress > div { height:100%; background:linear-gradient(90deg, var(--accent), #14b8a6); }
    .bars { display:grid; gap:10px; margin-top:12px; }
    .bar-row { display:grid; gap:8px; grid-template-columns:180px 1fr 70px; align-items:center; font-size:13px; }
    .bar-track { width:100%; height:12px; border-radius:999px; background:#ede3d5; overflow:hidden; }
    .bar-fill { height:100%; background:linear-gradient(90deg, var(--accent), #2dd4bf); }
    .bar-fill.orange { background:linear-gradient(90deg, var(--accent2), #fb923c); }
    .histogram { display:grid; grid-template-columns:repeat(auto-fit,minmax(48px,1fr)); align-items:end; gap:8px; min-height:180px; margin-top:12px; }
    .bin { display:grid; gap:6px; justify-items:center; font-size:11px; color:var(--muted); }
    .stick { width:100%; border-radius:10px 10px 2px 2px; background:linear-gradient(180deg,var(--accent),#2dd4bf); min-height:4px; }
    table { width:100%; border-collapse:collapse; font-size:13px; } th, td { border-bottom:1px solid var(--line); padding:10px 8px; text-align:left; vertical-align:top; }
    th { font-size:11px; text-transform:uppercase; color:var(--muted); letter-spacing:.08em; }
    .pill-row { display:flex; flex-wrap:wrap; gap:6px; }
    .pill { display:inline-flex; align-items:center; border-radius:999px; padding:4px 10px; font-size:12px; background:#efe4d5; color:var(--ink); }
    .step-list { display:grid; gap:8px; min-width:280px; }
    .step-row { display:grid; grid-template-columns:minmax(0,1fr) auto; gap:10px; padding:8px 10px; border:1px solid var(--line); border-radius:12px; background:#fff; }
    .step-row.slow { border-color:rgba(194,65,12,.35); background:rgba(251,146,60,.08); }
    .step-main { display:grid; gap:4px; }
    .step-message { font-weight:600; line-height:1.35; }
    .step-meta { font-size:11px; color:var(--muted); }
    .step-duration { align-self:start; white-space:nowrap; font-weight:700; color:var(--accent); }
    .step-duration.slow { color:var(--accent2); }
    .duration-badge { display:inline-flex; border-radius:999px; padding:4px 8px; background:rgba(15,118,110,.1); color:var(--accent); font-weight:700; }
    .duration-badge.slow { background:rgba(194,65,12,.12); color:var(--accent2); }
    .pager { display:flex; justify-content:flex-end; gap:8px; margin-top:12px; }
    .pager button { padding:8px 10px; border-radius:10px; }
    .pager .subtle { display:flex; align-items:center; }
    details.table-detail { min-width:280px; }
    details.table-detail > summary { cursor:pointer; color:var(--accent); font-weight:600; }
    details.table-detail[open] > summary { margin-bottom:10px; }
    .prose { line-height:1.6; white-space:pre-wrap; }
    pre { margin:0; white-space:pre-wrap; word-break:break-word; font-family:Consolas,monospace; font-size:12px; background:#fff; border:1px solid var(--line); border-radius:12px; padding:12px; max-height:280px; overflow:auto; }
    .empty { padding:18px; border:1px dashed var(--line); border-radius:16px; color:var(--muted); background:rgba(255,255,255,.5); text-align:center; }
    @media (max-width:1100px) { .page { grid-template-columns:1fr; } .sidebar { border-right:0; border-bottom:1px solid var(--line); } .row2 { grid-template-columns:1fr; } .span3, .span6, .span12 { grid-column:span 12; } .bar-row { grid-template-columns:1fr; } }
  </style>
</head>
<body>
  <div class=\"page\">
    <aside class=\"sidebar\">
      <section class=\"brand\"><h1>RAGFlow Performance Lab</h1><p>Uploads files, creates a dataset, measures parsing, benchmarks retrieval and chat, and reuses the same KB pipeline logs surfaced in the current RAGFlow UI.</p></section>
      <section class=\"panel\">
        <h2>New Run</h2>
        <form id=\"run-form\" class=\"form-grid\">
          <label>RAGFlow Base URL<input name=\"base_url\" value=\"http://127.0.0.1\" required /></label>
          <label>API Key<input name=\"api_key\" type=\"password\" required /></label>
          <label class=\"check\"><input name=\"verify_ssl\" type=\"checkbox\" checked />Verify SSL certificates</label>
          <div class=\"row2\"><label>HTTP Timeout (sec)<input name=\"http_timeout_sec\" type=\"number\" min=\"30\" max=\"1800\" value=\"180\" /></label><label>Parallel Runs<input name=\"parallel_runs\" type=\"number\" min=\"1\" max=\"50\" value=\"1\" /></label></div>
          <div class=\"row2\"><label>Dataset Prefix<input name=\"dataset_prefix\" value=\"perf-dataset\" /></label><label>Chat Prefix<input name=\"chat_prefix\" value=\"perf-chat\" /></label></div>
          <label>Files<input name=\"files\" type=\"file\" multiple required /></label>
          <label>Prompt Generation Settings
            <textarea name=\"prompt_options\">{
  \"prompts_per_document\": 3,
  \"shared_prompts\": 2
}</textarea>
          </label>
          <label>Dataset Options JSON
            <textarea name=\"dataset_options\">{
  \"chunk_method\": \"naive\",
  \"description\": \"Generated by performance load tester\",
  \"parser_config\": {
    \"layout_recognize\": \"DeepDOC\"
  }
}</textarea>
          </label>
          <label class=\"check\"><input name=\"enable_parsing\" type=\"checkbox\" checked />Enable parsing stage</label>
          <label>Parsing Config JSON
            <textarea name=\"parsing_options\">{
  \"poll_interval_sec\": 2,
  \"chunk_sample_size\": 3,
  \"collect_pipeline_logs\": true
}</textarea>
          </label>
          <label class=\"check\"><input name=\"enable_retrieval\" type=\"checkbox\" checked />Enable retrieval benchmark</label>
          <label>Retrieval Config JSON
            <textarea name=\"retrieval_options\">{
  \"concurrency\": 4,
  \"top_k\": 8,
  \"similarity_threshold\": 0.2,
  \"vector_similarity_weight\": 0.7,
  \"highlight\": true
}</textarea>
          </label>
          <label class=\"check\"><input name=\"enable_chat\" type=\"checkbox\" checked />Enable chat benchmark</label>
          <label>Chat Config JSON
            <span class=\"subtle\">Use the configured RAGFlow chat model identifier in `model_name`, for example `openai@OpenAI-API-Compatible` or `gpt-4o-mini@OpenAI-API-Compatible`. Leave it `null` to use the tenant default.</span>
            <textarea name=\"chat_options\">{
  \"concurrency\": 4,
  \"create\": {
    \"llm\": {
      \"model_name\": null
    }
  },
  \"completion\": {
    \"extra_body\": {
      \"reference\": true
    }
  }
}</textarea>
          </label>
          <label class=\"check\"><input name=\"enable_llm_summary\" type=\"checkbox\" checked />Generate executive results summary with the configured RAGFlow LLM</label>
          <label class=\"check\"><input name=\"cleanup_remote\" type=\"checkbox\" />Delete the remote dataset and chat after the run</label>
          <button type=\"submit\">Start Benchmark</button>
        </form>
      </section>
      <section class=\"panel\"><div style=\"display:grid;gap:10px;\"><div style=\"display:flex;justify-content:space-between;align-items:center;gap:10px;flex-wrap:wrap;\"><h2 style=\"margin:0;\">Runs</h2><div style=\"display:inline-flex;align-items:center;gap:8px;color:var(--muted);white-space:nowrap;\">Auto Refresh (sec)<input id=\"auto-refresh-sec\" type=\"number\" min=\"0\" max=\"3600\" value=\"0\" style=\"width:92px;\" /></div></div><div style=\"display:flex;gap:8px;justify-content:flex-end;flex-wrap:wrap;\"><button class=\"secondary\" id=\"refresh-runs\" type=\"button\">Refresh</button><button class=\"secondary\" id=\"reset-runs\" type=\"button\">Reset</button></div></div><div id=\"run-list\" class=\"run-list\"></div></section>
    </aside>
    <main class=\"content\"><div id=\"run-view\" class=\"empty\">Start a run or select a previous benchmark from the left.</div></main>
  </div>
<script>
const state = { activeRunId: null, activeBatchTabs: {}, tablePages: {}, chartPages: {}, pollHandle: null, autoRefreshHandle: null, runs: [] };
const BASE_PATH = '__BASE_PATH__';
const FORM_STORAGE_KEY = 'ragflow-performance-form-v1';
const ACTIVE_RUN_STORAGE_KEY = 'ragflow-performance-active-run-v1';
const AUTO_REFRESH_STORAGE_KEY = 'ragflow-performance-auto-refresh-v1';
const FORM_FIELDS = ['base_url', 'api_key', 'verify_ssl', 'http_timeout_sec', 'parallel_runs', 'dataset_prefix', 'chat_prefix', 'prompt_options', 'dataset_options', 'enable_parsing', 'parsing_options', 'enable_retrieval', 'retrieval_options', 'enable_chat', 'chat_options', 'enable_llm_summary', 'cleanup_remote'];
function statusBadge(status) { const normalized = (status || 'queued').toLowerCase(); return `<span class=\"badge ${normalized}\">${normalized}</span>`; }
function number(value, digits = 2) { if (value === null || value === undefined || Number.isNaN(Number(value))) return '-'; return Number(value).toFixed(digits); }
function formatTs(value) { if (!value) return '-'; const date = new Date(value); return Number.isNaN(date.getTime()) ? value : date.toLocaleString(); }
function metricCard(label, value, hint = '') { return `<div class=\"card span3\"><div class=\"metric\"><div class=\"label\">${label}</div><div class=\"value\">${value}</div><div class=\"subtle\">${hint}</div></div></div>`; }
function renderBars(rows, formatter = (value) => value, orange = false) { if (!rows || !rows.length) return `<div class=\"empty\">No chart data available.</div>`; const max = Math.max(...rows.map((row) => Number(row.value) || 0), 1); return `<div class=\"bars\">${rows.map((row) => `<div class=\"bar-row\"><div title=\"${row.label}\">${row.label}</div><div class=\"bar-track\"><div class=\"bar-fill ${orange ? 'orange' : ''}\" style=\"width:${Math.max(4, (Number(row.value) || 0) / max * 100)}%\"></div></div><div>${formatter(row.value)}</div></div>`).join('')}</div>`; }
function renderHistogram(data) { if (!data || !data.length) return `<div class=\"empty\">No latency data available.</div>`; const max = Math.max(...data.map((bin) => bin.count), 1); return `<div class=\"histogram\">${data.map((bin) => `<div class=\"bin\"><div>${bin.count}</div><div class=\"stick\" style=\"height:${Math.max(8, bin.count / max * 140)}px\"></div><div>${bin.label}</div></div>`).join('')}</div>`; }
function renderTable(headers, rows) { if (!rows || !rows.length) return `<div class=\"empty\">No rows available.</div>`; return `<table><thead><tr>${headers.map((header) => `<th>${header}</th>`).join('')}</tr></thead><tbody>${rows.map((row) => `<tr>${row.map((cell) => `<td>${cell}</td>`).join('')}</tr>`).join('')}</tbody></table>`; }
function renderPaginatedTable(tableId, headers, rows, pageSize = 10) { if (!rows || !rows.length) return `<div class=\"empty\">No rows available.</div>`; const totalPages = Math.max(1, Math.ceil(rows.length / pageSize)); const currentPage = Math.min(totalPages, Math.max(1, Number(state.tablePages[tableId] || 1))); state.tablePages[tableId] = currentPage; const start = (currentPage - 1) * pageSize; const pagedRows = rows.slice(start, start + pageSize); const pager = totalPages > 1 ? `<div class=\"pager\"><button class=\"secondary\" type=\"button\" data-table-page=\"${Math.max(1, currentPage - 1)}\" data-table-id=\"${tableId}\" ${currentPage === 1 ? 'disabled' : ''}>Prev</button><div class=\"subtle\">Page ${currentPage} / ${totalPages}</div><button class=\"secondary\" type=\"button\" data-table-page=\"${Math.min(totalPages, currentPage + 1)}\" data-table-id=\"${tableId}\" ${currentPage === totalPages ? 'disabled' : ''}>Next</button></div>` : ''; return `${renderTable(headers, pagedRows)}${pager}`; }
function renderPaginatedBars(chartId, rows, pageSize = 12, formatter = (value) => value, orange = false) { if (!rows || !rows.length) return `<div class=\"empty\">No chart data available.</div>`; const totalPages = Math.max(1, Math.ceil(rows.length / pageSize)); const currentPage = Math.min(totalPages, Math.max(1, Number(state.chartPages[chartId] || 1))); state.chartPages[chartId] = currentPage; const start = (currentPage - 1) * pageSize; const pagedRows = rows.slice(start, start + pageSize); const pager = totalPages > 1 ? `<div class=\"pager\"><button class=\"secondary\" type=\"button\" data-chart-page=\"${Math.max(1, currentPage - 1)}\" data-chart-id=\"${chartId}\" ${currentPage === 1 ? 'disabled' : ''}>Prev</button><div class=\"subtle\">Page ${currentPage} / ${totalPages}</div><button class=\"secondary\" type=\"button\" data-chart-page=\"${Math.min(totalPages, currentPage + 1)}\" data-chart-id=\"${chartId}\" ${currentPage === totalPages ? 'disabled' : ''}>Next</button></div>` : ''; return `${renderBars(pagedRows, formatter, orange)}${pager}`; }
function escapeHtml(value) { return String(value ?? '').replaceAll('&', '&amp;').replaceAll('<', '&lt;').replaceAll('>', '&gt;').replaceAll('\"', '&quot;').replaceAll(\"'\", '&#39;'); }
function renderDurationBadge(value, threshold) { const duration = Number(value) || 0; if (!duration) return '-'; const slow = threshold > 0 && duration >= threshold; return `<span class=\"duration-badge ${slow ? 'slow' : ''}\">${number(duration, 2)} s</span>`; }
function renderPipelinePath(item) { const path = Array.isArray(item.pipeline_path) ? item.pipeline_path.filter(Boolean) : []; if (!path.length) return `<span class=\"subtle\">No pipeline path returned by this log.</span>`; return `<div class=\"pill-row\">${path.map((part) => `<span class=\"pill\">${escapeHtml(part)}</span>`).join('')}</div>`; }
function renderProgressSteps(item) { const steps = Array.isArray(item.progress_steps) ? item.progress_steps : []; if (!steps.length) { const fallback = item.progress_msg ? escapeHtml(item.progress_msg) : 'No progress steps returned.'; return `<div class=\"subtle\">${fallback}</div>`; } return `<div class=\"step-list\">${steps.map((step, index) => { const slow = Boolean(step.is_slow); const duration = Number(step.duration_sec); const timestamp = step.timestamp ? escapeHtml(step.timestamp) : `step ${index + 1}`; const meta = duration > 0 ? `${timestamp} • next step in ${number(duration, 2)} s` : timestamp; return `<div class=\"step-row ${slow ? 'slow' : ''}\"><div class=\"step-main\"><div class=\"step-message\">${escapeHtml(step.message || 'step')}</div><div class=\"step-meta\">${meta}</div></div><div class=\"step-duration ${slow ? 'slow' : ''}\">${duration > 0 ? `${number(duration, 2)} s` : '-'}</div></div>`; }).join('')}</div>`; }
function renderPipelineDetails(item) { return `<details class=\"table-detail\"><summary>Show details</summary><div style=\"display:grid;gap:10px;\"><div><div class=\"subtle\" style=\"margin-bottom:6px;\">Pipeline path</div>${renderPipelinePath(item)}</div><div><div class=\"subtle\" style=\"margin-bottom:6px;\">Progress steps</div>${renderProgressSteps(item)}</div>${item.progress_msg ? `<div><div class=\"subtle\" style=\"margin-bottom:6px;\">Raw message</div><pre>${escapeHtml(item.progress_msg)}</pre></div>` : ''}</div></details>`; }
function persistFormState() { try { const form = document.getElementById('run-form'); const payload = {}; FORM_FIELDS.forEach((name) => { const field = form.elements.namedItem(name); if (!field) return; if (field.type === 'checkbox') payload[name] = field.checked; else payload[name] = field.value; }); localStorage.setItem(FORM_STORAGE_KEY, JSON.stringify(payload)); } catch (error) { console.warn('Unable to persist form state.', error); } }
function restoreFormState() { try { const raw = localStorage.getItem(FORM_STORAGE_KEY); if (!raw) return; const payload = JSON.parse(raw); const form = document.getElementById('run-form'); FORM_FIELDS.forEach((name) => { if (!(name in payload)) return; const field = form.elements.namedItem(name); if (!field) return; if (field.type === 'checkbox') field.checked = Boolean(payload[name]); else if (typeof payload[name] === 'string' || typeof payload[name] === 'number') field.value = payload[name]; }); } catch (error) { console.warn('Unable to restore form state.', error); } }
function persistActiveRun() { try { if (state.activeRunId) localStorage.setItem(ACTIVE_RUN_STORAGE_KEY, state.activeRunId); else localStorage.removeItem(ACTIVE_RUN_STORAGE_KEY); } catch (error) { console.warn('Unable to persist active run.', error); } }
function restoreActiveRun() { try { return localStorage.getItem(ACTIVE_RUN_STORAGE_KEY); } catch (error) { console.warn('Unable to restore active run.', error); return null; } }
function persistAutoRefreshSetting() { try { const field = document.getElementById('auto-refresh-sec'); localStorage.setItem(AUTO_REFRESH_STORAGE_KEY, String(Math.max(0, Number(field?.value || 0) || 0))); } catch (error) { console.warn('Unable to persist auto refresh setting.', error); } }
function restoreAutoRefreshSetting() { try { const saved = localStorage.getItem(AUTO_REFRESH_STORAGE_KEY); const field = document.getElementById('auto-refresh-sec'); if (field && saved !== null) field.value = String(Math.max(0, Number(saved) || 0)); } catch (error) { console.warn('Unable to restore auto refresh setting.', error); } }
function configureAutoRefresh() { if (state.autoRefreshHandle) { window.clearInterval(state.autoRefreshHandle); state.autoRefreshHandle = null; } const seconds = Math.max(0, Number(document.getElementById('auto-refresh-sec')?.value || 0) || 0); if (seconds > 0) state.autoRefreshHandle = window.setInterval(() => window.location.reload(), seconds * 1000); }
function stopAutoRefreshIfIdle() { const hasLiveRuns = state.runs.some((run) => ['queued', 'running'].includes((run.status || '').toLowerCase())); if (hasLiveRuns) return; const field = document.getElementById('auto-refresh-sec'); if (!field) return; if (Number(field.value || 0) <= 0) return; field.value = '0'; persistAutoRefreshSetting(); configureAutoRefresh(); }
function upsertRun(run) { const index = state.runs.findIndex((item) => item.id === run.id); if (index >= 0) state.runs[index] = run; else state.runs.unshift(run); }
function getRun(runId) { return state.runs.find((item) => item.id === runId) || null; }
function getBatchRuns(batchId) { if (!batchId) return []; return state.runs.filter((item) => item.config?.batch?.batch_id === batchId).sort((a, b) => Number(a.config?.batch?.run_index || 0) - Number(b.config?.batch?.run_index || 0)); }
function summarizeBatchRuns(batchRuns) { const average = (values) => values.length ? values.reduce((sum, value) => sum + value, 0) / values.length : 0; const valuesFor = (field) => batchRuns.map((run) => Number(run.summary?.[field] || 0)).filter((value) => value > 0); return { total_runs: batchRuns.length, completed_runs: batchRuns.filter((run) => (run.status || '').toLowerCase() === 'completed').length, failed_runs: batchRuns.filter((run) => (run.status || '').toLowerCase() === 'failed').length, avg_parse_wall_time_sec: average(valuesFor('parse_wall_time_sec')), avg_retrieval_p95_ms: average(valuesFor('retrieval_p95_ms')), max_retrieval_p95_ms: valuesFor('retrieval_p95_ms').length ? Math.max(...valuesFor('retrieval_p95_ms')) : 0, avg_chat_p95_ms: average(valuesFor('chat_p95_ms')), max_chat_p95_ms: valuesFor('chat_p95_ms').length ? Math.max(...valuesFor('chat_p95_ms')) : 0, avg_chat_error_rate: average(batchRuns.map((run) => Number(run.summary?.chat_error_rate || 0))), avg_chat_tokens_per_sec: average(valuesFor('chat_tokens_per_sec')) }; }
function renderBatchTabs(batchRuns, activeTab) { if (batchRuns.length <= 1) return ''; return `<div class=\"tabs\"><button class=\"tab ${activeTab === 'overview' ? 'active' : ''}\" data-batch-tab=\"overview\" type=\"button\">Batch Overview</button>${batchRuns.map((run) => `<button class=\"tab ${activeTab === run.id ? 'active' : ''}\" data-batch-tab=\"${run.id}\" type=\"button\">Run ${run.config?.batch?.run_index || '?'}</button>`).join('')}</div>`; }
function renderBatchOverview(run, batchRuns) { const batch = run.config?.batch || {}; const batchSummary = batchRuns.find((item) => item.batch_summary?.aggregate || item.batch_summary?.llm_assessment)?.batch_summary || {}; const aggregate = batchSummary.aggregate || summarizeBatchRuns(batchRuns); const llm = batchSummary.llm_assessment?.content || ''; const canDownload = (aggregate.completed_runs || 0) + (aggregate.failed_runs || 0) === (aggregate.total_runs || batchRuns.length) && (aggregate.total_runs || batchRuns.length) > 0; const reportHref = batch.batch_id ? `${BASE_PATH}/api/batches/${encodeURIComponent(batch.batch_id)}/report.docx` : '#'; const rows = batchRuns.map((item) => [String(item.config?.batch?.run_index || '-'), statusBadge(item.status), escapeHtml(item.dataset?.name || item.config?.dataset_prefix || '-'), `${number(item.summary?.parse_wall_time_sec || 0)} s`, `${number(item.summary?.retrieval_p95_ms || 0)} ms`, `${number(item.summary?.chat_p95_ms || 0)} ms`, `${number((item.summary?.chat_error_rate || 0) * 100)}%`, item.error ? `<span style=\"color:var(--bad);\">${escapeHtml(item.error)}</span>` : '-']); return `<div style=\"display:flex;justify-content:space-between;align-items:flex-start;gap:20px;margin-bottom:18px;\"><div><div style=\"display:flex;gap:10px;align-items:center;flex-wrap:wrap;\"><h2 style=\"margin:0;font-size:34px;letter-spacing:-.06em;\">${escapeHtml(run.config?.dataset_prefix || 'Batch')} batch</h2>${statusBadge(aggregate.failed_runs ? 'failed' : (aggregate.completed_runs === aggregate.total_runs ? 'completed' : 'running'))}</div><div class=\"subtle\" style=\"margin-top:6px;\">${escapeHtml(batch.batch_id || '-')} • ${batch.parallel_count || batch.run_count || batchRuns.length} runs • host ${escapeHtml(run.config?.base_url || '-')}</div></div><div style=\"display:grid;justify-items:end;gap:8px;\"><div class=\"subtle\" style=\"text-align:right;\"><div>First start: ${formatTs(batchRuns[0]?.started_at || batchRuns[0]?.created_at)}</div><div>Latest completion: ${formatTs(batchRuns[batchRuns.length - 1]?.completed_at)}</div></div>${canDownload ? `<a class=\"secondary\" style=\"text-decoration:none;display:inline-flex;align-items:center;justify-content:center;\" href=\"${reportHref}\">Download Batch Word Report</a>` : `<div class=\"subtle\">Batch report is available after all runs finish.</div>`}</div></div><section class=\"grid\">${metricCard('Runs', aggregate.total_runs || batchRuns.length, `${aggregate.completed_runs || 0} completed • ${aggregate.failed_runs || 0} failed`)}${metricCard('Avg Parse Wall', `${number(aggregate.avg_parse_wall_time_sec || 0)} s`, 'wall time per run')}${metricCard('Avg Retrieval P95', `${number(aggregate.avg_retrieval_p95_ms || 0)} ms`, `max ${number(aggregate.max_retrieval_p95_ms || 0)} ms`)}${metricCard('Avg Chat P95', `${number(aggregate.avg_chat_p95_ms || 0)} ms`, `max ${number(aggregate.max_chat_p95_ms || 0)} ms`)}${metricCard('Avg Chat Error Rate', `${number((aggregate.avg_chat_error_rate || 0) * 100)}%`, `${number(aggregate.avg_chat_tokens_per_sec || 0)} tok/s`)}<div class=\"card span12\"><h3>Batch Executive Summary</h3><div class=\"subtle\">${llm ? 'Generated by the configured RAGFlow chat model after the batch finished.' : 'This appears after every run in the batch reaches a terminal state.'}</div>${llm ? `<div class=\"prose\">${escapeHtml(llm)}</div>` : ''}</div><div class=\"card span12\"><h3>Batch Run Table</h3>${renderPaginatedTable(`batch-${batch.batch_id || 'default'}-runs`, ['run','status','dataset','parse wall','retrieval p95','chat p95','chat errors','error'], rows, 10)}</div></section>`; }
function attachBatchTabHandlers(container, run, batchRuns) { const batch = run.config?.batch || {}; if ((batch.run_count || 1) <= 1 || !batch.batch_id) return; container.querySelectorAll('[data-batch-tab]').forEach((element) => element.addEventListener('click', async () => { const tab = element.getAttribute('data-batch-tab'); state.activeBatchTabs[batch.batch_id] = tab; if (tab && tab !== 'overview') { state.activeRunId = tab; persistActiveRun(); const selected = getRun(tab); renderRunView(selected || run); if (selected && ['queued', 'running'].includes((selected.status || '').toLowerCase())) startPolling(tab); else stopPolling(); } else { const liveRun = batchRuns.find((item) => ['queued', 'running'].includes((item.status || '').toLowerCase())); state.activeRunId = liveRun?.id || run.id; persistActiveRun(); renderRunView(run); if (liveRun) startPolling(liveRun.id); else stopPolling(); } renderRunList(); })); }
function attachPaginationHandlers(container) { container.querySelectorAll('[data-table-id][data-table-page]').forEach((element) => element.addEventListener('click', () => { const tableId = element.getAttribute('data-table-id'); const page = Number(element.getAttribute('data-table-page') || 1); state.tablePages[tableId] = page; const active = getRun(state.activeRunId); if (active) renderRunView(active); })); }
function attachChartPaginationHandlers(container) { container.querySelectorAll('[data-chart-id][data-chart-page]').forEach((element) => element.addEventListener('click', () => { const chartId = element.getAttribute('data-chart-id'); const page = Number(element.getAttribute('data-chart-page') || 1); state.chartPages[chartId] = page; const active = getRun(state.activeRunId); if (active) renderRunView(active); })); }
async function fetchRuns() { const response = await fetch(`${BASE_PATH}/api/runs`); const payload = await response.json(); state.runs = payload.runs || []; renderRunList(); stopAutoRefreshIfIdle(); const targetRunId = state.activeRunId || restoreActiveRun() || state.runs.find((run) => ['queued', 'running'].includes((run.status || '').toLowerCase()))?.id; if (targetRunId) { const active = getRun(targetRunId); if (active) { state.activeRunId = targetRunId; persistActiveRun(); renderRunView(active); if (['queued', 'running'].includes((active.status || '').toLowerCase())) startPolling(targetRunId); else stopPolling(); } else { state.activeRunId = null; persistActiveRun(); stopPolling(); document.getElementById('run-view').innerHTML = `<div class=\"empty\">Start a run or select a previous benchmark from the left.</div>`; } } }
async function fetchRun(runId, options = {}) { if (!runId) return; const response = await fetch(`${BASE_PATH}/api/runs/${runId}`); const payload = await response.json(); if (!response.ok || payload.error) return; upsertRun(payload); state.activeRunId = runId; persistActiveRun(); renderRunView(payload); if (!options.silent) renderRunList(); stopAutoRefreshIfIdle(); const live = ['queued', 'running'].includes((payload.status || '').toLowerCase()); if (live) startPolling(runId); else stopPolling(); }
function startPolling(runId) { stopPolling(); state.pollHandle = window.setInterval(() => fetchRun(runId, { silent: true }), 2500); }
function stopPolling() { if (state.pollHandle) { window.clearInterval(state.pollHandle); state.pollHandle = null; } }
function renderRunList() { const container = document.getElementById('run-list'); if (!state.runs.length) { container.innerHTML = `<div class=\"empty\">No runs yet.</div>`; return; } container.innerHTML = state.runs.map((run) => { const stages = []; if (run.config?.stages?.parsing?.enabled) stages.push('parse'); if (run.config?.stages?.retrieval?.enabled) stages.push('retrieval'); if (run.config?.stages?.chat?.enabled) stages.push('chat'); const batch = run.config?.batch || {}; const batchLabel = batch.run_count > 1 ? ` • run ${batch.run_index}/${batch.run_count}` : ''; return `<div class=\"run-item ${run.id === state.activeRunId ? 'active' : ''}\" data-run-id=\"${run.id}\"><div style=\"display:flex;justify-content:space-between;align-items:flex-start;gap:10px;\"><div><div style=\"font-weight:700;line-height:1.25;\">${escapeHtml(run.dataset?.name || run.config?.dataset_prefix || 'pending run')}</div><div class=\"subtle\">${run.summary?.prompts || 0} prompts • ${(stages.join(' + ') || 'upload-only')}${batchLabel}</div></div>${statusBadge(run.status)}</div><div class=\"meta\" style=\"margin-top:8px;\"><span>${formatTs(run.started_at || run.created_at)}</span><span>${escapeHtml(run.phase || '-')}</span></div></div>`; }).join(''); container.querySelectorAll('.run-item').forEach((element) => element.addEventListener('click', () => { const run = getRun(element.dataset.runId); const batchId = run?.config?.batch?.batch_id; if (batchId && (run?.config?.batch?.run_count || 1) > 1) state.activeBatchTabs[batchId] = run.id; state.activeRunId = element.dataset.runId; persistActiveRun(); fetchRun(element.dataset.runId); })); }
function renderRunView(run) {
  upsertRun(run);
  const container = document.getElementById('run-view');
  const batchRuns = getBatchRuns(run.config?.batch?.batch_id);
  const hasBatch = (run.config?.batch?.run_count || 1) > 1 && batchRuns.length > 0;
  if (hasBatch && !state.activeBatchTabs[run.config.batch.batch_id]) state.activeBatchTabs[run.config.batch.batch_id] = 'overview';
  const activeBatchTab = hasBatch ? state.activeBatchTabs[run.config.batch.batch_id] || 'overview' : run.id;
  if (hasBatch && activeBatchTab === 'overview') {
    container.innerHTML = `${renderBatchTabs(batchRuns, activeBatchTab)}${renderBatchOverview(run, batchRuns)}`;
    attachBatchTabHandlers(container, run, batchRuns);
    attachPaginationHandlers(container);
    return;
  }
  if (hasBatch && activeBatchTab && activeBatchTab !== run.id) {
    const selectedRun = getRun(activeBatchTab);
    if (selectedRun) run = selectedRun;
  }
  const parse = run.parse || {}; const retrieval = run.retrieval || { summary: {}, results: [] }; const chat = run.chat_results || { summary: {}, results: [] }; const live = run.live || {}; const prompts = run.prompts || []; const events = run.events || []; const timeline = Array.isArray(run.timeline) ? run.timeline : [];
  const stages = run.config?.stages || {};
  const batch = run.config?.batch || {};
  const batchLabel = batch.run_count > 1 ? ` • run ${batch.run_index}/${batch.run_count}` : '';
  const parsingEnabled = Boolean(stages.parsing?.enabled);
  const retrievalEnabled = Boolean(stages.retrieval?.enabled);
  const chatEnabled = Boolean(stages.chat?.enabled);
  const metricValue = (enabled, value, digits = 2, suffix = '') => enabled ? `${number(value, digits)}${suffix}` : 'skipped';
  const metricHint = (enabled, hint) => enabled ? hint : 'stage disabled';
  const parseBars = (parse.documents || []).map((doc) => ({ label: doc.name, value: Number(doc.process_duration || 0) })).sort((a,b) => b.value - a.value);
  const stageDurations = timeline.map((stage) => Number(stage.duration_sec || 0)).filter((value) => value > 0).sort((a, b) => a - b);
  const stageSlowThreshold = stageDurations.length ? stageDurations[Math.floor((stageDurations.length - 1) * 0.75)] : 0;
  const stageBars = timeline.filter((stage) => Number(stage.duration_sec || 0) > 0).map((stage) => ({ label: stage.label || stage.key || 'stage', value: Number(stage.duration_sec || 0) })).sort((a,b) => b.value - a.value);
  const retrievalRows = (retrieval.results || []).map((item) => [item.ok ? 'ok' : `<span style=\"color:var(--bad);\">error</span>`, item.kind, item.document_name || '-', item.prompt, `${number(item.latency_ms, 1)} ms`, String(item.chunk_count), item.top_documents?.join(', ') || '-']);
  const chatRows = (chat.results || []).map((item) => [item.ok ? 'ok' : `<span style=\"color:var(--bad);\">error</span>`, item.kind, item.document_name || '-', item.prompt, `${number(item.latency_ms, 1)} ms`, `${item.total_tokens || 0}`, `${item.reference_count || 0}`, item.referenced_documents?.join(', ') || '-']);
  const stageRows = timeline.map((stage) => [escapeHtml(stage.label || stage.key || '-'), statusBadge(stage.status || 'queued'), renderDurationBadge(stage.duration_sec, stageSlowThreshold), formatTs(stage.started_at), formatTs(stage.completed_at)]);
  const summaryCard = run.analysis?.llm_assessment?.content ? `<div class=\"card span12\"><h3>Executive Summary</h3><div class=\"subtle\">Generated by the configured RAGFlow chat model.</div><div class=\"prose\">${escapeHtml(run.analysis.llm_assessment.content)}</div></div>` : '';
  container.innerHTML = `
    ${renderBatchTabs(batchRuns, activeBatchTab)}
    <div style=\"display:flex;justify-content:space-between;align-items:flex-start;gap:20px;margin-bottom:18px;\"><div><div style=\"display:flex;gap:10px;align-items:center;flex-wrap:wrap;\"><h2 style=\"margin:0;font-size:34px;letter-spacing:-.06em;\">${run.dataset?.name || 'Pending run'}</h2>${statusBadge(run.status)}</div><div class=\"subtle\" style=\"margin-top:6px;\">${run.dataset?.id || '-'} • ${run.phase || '-'}${batchLabel} • ${run.error ? `<span style=\"color:var(--bad);\">${run.error}</span>` : 'no fatal error'}</div></div><div class=\"subtle\" style=\"text-align:right;\"><div>Started: ${formatTs(run.started_at || run.created_at)}</div><div>Completed: ${formatTs(run.completed_at)}</div><div>Host: ${run.config?.base_url || '-'}</div></div></div>
    <section class=\"card\" style=\"margin-bottom:16px;\"><div style=\"display:flex;justify-content:space-between;gap:10px;align-items:center;\"><div><h3 style=\"margin:0 0 6px;\">Live Progress</h3><div class=\"subtle\">${live.parse_completed || 0}/${live.parse_total || 0} parsed • ${live.retrieval_completed || 0}/${live.retrieval_total || 0} retrievals • ${live.chat_completed || 0}/${live.chat_total || 0} chats</div></div><div style=\"font-size:30px;font-weight:800;letter-spacing:-.05em;\">${Number(run.progress || 0)}%</div></div><div class=\"progress\"><div style=\"width:${Math.max(2, Number(run.progress || 0))}%\"></div></div><div class=\"subtle\">${events.length ? events[events.length - 1].message : 'Waiting for updates.'}</div></section>
    <section class=\"grid\">
      ${metricCard('Documents', run.summary?.documents || parse.total_documents || 0, parsingEnabled ? `${parse.total_chunks || 0} chunks • ${parse.total_tokens || 0} tokens indexed` : 'uploaded files only')}
      ${metricCard('Prompts', run.summary?.prompts || prompts.length || 0, `${run.config?.prompts_per_document || 0}/doc + ${run.config?.shared_prompts || 0} shared`)}
      ${metricCard('Parse P95', metricValue(parsingEnabled, parse.p95_parse_duration_sec || 0, 2, ' s'), metricHint(parsingEnabled, `wall ${number(run.summary?.parse_wall_time_sec || parse.wall_time_sec || 0)} s`))}
      ${metricCard('Chat P95', metricValue(chatEnabled, chat.summary?.p95_latency_ms || 0, 2, ' ms'), metricHint(chatEnabled, `${number(chat.summary?.throughput_rps || 0)} req/s`))}
      ${metricCard('Retrieval P95', metricValue(retrievalEnabled, retrieval.summary?.p95_latency_ms || 0, 2, ' ms'), metricHint(retrievalEnabled, `${number(retrieval.summary?.throughput_rps || 0)} req/s`))}
      ${metricCard('Chat Error Rate', chatEnabled ? `${number((chat.summary?.error_rate || 0) * 100)}%` : 'skipped', metricHint(chatEnabled, `${chat.summary?.errors || 0} failed calls`))}
      ${metricCard('Token Throughput', metricValue(chatEnabled, chat.summary?.tokens_per_sec || 0, 2, ' tok/s'), metricHint(chatEnabled, `${chat.summary?.total_tokens || 0} total tokens`))}
      ${metricCard('Parse Token Rate', metricValue(parsingEnabled, parse.token_throughput_per_sec || 0, 2, ' tok/s'), metricHint(parsingEnabled, `${parse.status_breakdown ? Object.entries(parse.status_breakdown).map(([k,v]) => `${k}:${v}`).join(' • ') : '-'}`))}
      ${summaryCard}
      <div class=\"card span6\"><h3>Parse Duration by Document</h3><div class=\"subtle\">Total parse time per document from the final document status API. All documents are available through pagination.</div>${renderPaginatedBars(`parse-bars-${run.id}`, parseBars, 12, (value) => `${number(value)} s`)}</div>
      <div class=\"card span6\"><h3>Execution Stage Durations</h3><div class=\"subtle\">Measured directly by this app from the run flow: provisioning, upload, parsing, prompt generation, retrieval, assistant setup, chat, summary, and cleanup.${stageSlowThreshold > 0 ? ` Stages at or above ${number(stageSlowThreshold, 2)} s are treated as slow.` : ''}</div>${renderPaginatedBars(`stage-bars-${run.id}`, stageBars, 12, (value) => `${number(value)} s`, true)}</div>
      <div class=\"card span6\"><h3>Retrieval Latency Histogram</h3><div class=\"subtle\">Official /api/v1/retrieval benchmark on the generated prompts.</div>${renderHistogram(retrieval.summary?.latency_histogram || [])}</div>
      <div class=\"card span6\"><h3>Chat Latency Histogram</h3><div class=\"subtle\">Official OpenAI-compatible chat completions.</div>${renderHistogram(chat.summary?.latency_histogram || [])}</div>
      <div class=\"card span12\"><h3>Execution Stage Timeline</h3><div class=\"subtle\">This is the app-controlled workflow, not RagFlow KB pipeline metadata.</div>${renderPaginatedTable(`timeline-${run.id}`, ['stage','status','duration','started','completed'], stageRows, 10)}</div>
      <div class=\"card span12\"><h3>Retrieval Sample Results</h3>${renderPaginatedTable(`retrieval-${run.id}`, ['status','kind','document','prompt','latency','chunks','top docs'], retrievalRows, 10)}</div>
      <div class=\"card span12\"><h3>Chat Sample Results</h3>${renderPaginatedTable(`chat-${run.id}`, ['status','kind','document','prompt','latency','tokens','refs','referenced docs'], chatRows, 10)}</div>
      <div class=\"card span6\"><h3>Generated Prompts</h3><pre>${JSON.stringify(prompts, null, 2)}</pre></div>
      <div class=\"card span6\"><h3>Event Log</h3><pre>${events.map((event) => `[${formatTs(event.ts)}] ${event.level.toUpperCase()} ${event.message}`).join('\\n')}</pre></div>
    </section>`;
  attachBatchTabHandlers(container, run, batchRuns);
  attachPaginationHandlers(container);
  attachChartPaginationHandlers(container);
}
document.getElementById('refresh-runs').addEventListener('click', fetchRuns);
document.getElementById('reset-runs').addEventListener('click', async () => { if (!window.confirm('Clear all stored run results from this app? Active configs in local storage will be kept.')) return; const response = await fetch(`${BASE_PATH}/api/runs/reset`, { method: 'POST' }); const payload = await response.json(); if (!response.ok || payload.error) { alert(payload.error || 'Unable to reset runs.'); return; } state.activeRunId = null; state.activeBatchTabs = {}; persistActiveRun(); stopPolling(); document.getElementById('run-view').innerHTML = `<div class=\"empty\">Start a run or select a previous benchmark from the left.</div>`; await fetchRuns(); });
document.getElementById('run-form').addEventListener('input', persistFormState);
document.getElementById('run-form').addEventListener('change', persistFormState);
document.getElementById('auto-refresh-sec').addEventListener('input', () => { persistAutoRefreshSetting(); configureAutoRefresh(); });
document.getElementById('auto-refresh-sec').addEventListener('change', () => { persistAutoRefreshSetting(); configureAutoRefresh(); });
document.getElementById('run-form').addEventListener('submit', async (event) => { event.preventDefault(); const formData = new FormData(event.currentTarget); const response = await fetch(`${BASE_PATH}/api/start`, { method: 'POST', body: formData }); const payload = await response.json(); if (!response.ok || payload.error) { alert(payload.error || 'Run creation failed.'); return; } const targetRunId = payload.run_id || (payload.run_ids && payload.run_ids[0]); persistFormState(); await fetchRuns(); if (targetRunId) await fetchRun(targetRunId); });
restoreFormState();
restoreAutoRefreshSetting();
configureAutoRefresh();
fetchRuns();
</script>
</body>
</html>"""

def render_index_html() -> str:
    return INDEX_HTML.replace("__BASE_PATH__", APP_BASE_PATH)


@app.get(route_path("/"))
async def index() -> str:
    return render_index_html()


if APP_BASE_PATH:
    app.add_url_rule(f"{route_path('/')}/", view_func=index)


@app.get(route_path("/api/runs"))
async def list_runs() -> Any:
    with RUNS_LOCK:
        runs = sorted((json_clone(run) for run in RUNS.values()), key=lambda item: item.get("created_at", ""), reverse=True)
    return jsonify({"runs": runs})


@app.get(route_path("/api/runs/<run_id>"))
async def get_run(run_id: str) -> Any:
    with RUNS_LOCK:
        if run_id not in RUNS:
            return jsonify({"error": "Run not found"}), 404
    return jsonify(build_run_response(run_id))


@app.get(route_path("/api/runs/<run_id>/report.docx"))
async def download_run_report(run_id: str) -> Any:
    with RUNS_LOCK:
        if run_id not in RUNS:
            return jsonify({"error": "Run not found"}), 404
        run = json_clone(RUNS[run_id])

    report_bytes = build_run_report(run)
    filename = f"ragflow-performance-{run_id}.docx"
    return Response(
        report_bytes,
        headers={
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(len(report_bytes)),
        },
    )


@app.get(route_path("/api/batches/<batch_id>/report.docx"))
async def download_batch_report(batch_id: str) -> Any:
    with RUNS_LOCK:
        batch_runs = [json_clone(run) for run in RUNS.values() if run.get("config", {}).get("batch", {}).get("batch_id") == batch_id]
    if not batch_runs:
        return jsonify({"error": "Batch not found"}), 404

    report_bytes = build_batch_report(batch_id, batch_runs)
    filename = f"ragflow-performance-batch-{batch_id}.docx"
    return Response(
        report_bytes,
        headers={
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(len(report_bytes)),
        },
    )


@app.post(route_path("/api/runs/reset"))
async def reset_runs() -> Any:
    with RUNS_LOCK:
        active = [run_id for run_id, run in RUNS.items() if (run.get("status") or "").lower() in {"queued", "running"} and run_id in RUN_SECRETS]
        if active:
            return jsonify({"error": "Cannot reset while runs are still queued or running."}), 409
        removed = list(RUNS.keys())
        RUNS.clear()
        RUN_SECRETS.clear()
    with BATCH_LIMITERS_LOCK:
        BATCH_LIMITERS.clear()
        BATCH_ANALYSIS_STATE.clear()
    for json_file in RUNS_DIR.glob("*.json"):
        try:
            json_file.unlink()
        except OSError:
            continue
    return jsonify({"ok": True, "cleared_runs": len(removed)})


@app.post(route_path("/api/start"))
async def start_run() -> Any:
    form = await request.form
    files = await request.files
    uploaded_files = files.getlist("files")
    if not uploaded_files:
        return jsonify({"error": "Please upload at least one file."}), 400

    try:
        base_url = form.get("base_url", "").strip()
        api_key = form.get("api_key", "").strip()
        normalize_root_url(base_url)
        auth_header(api_key)
        prompt_options = parse_json_object(form.get("prompt_options"), "Prompt generation settings", {"prompts_per_document": 3, "shared_prompts": 2})
        dataset_options = parse_json_object(
            form.get("dataset_options"),
            "Dataset options",
            {
                "chunk_method": "naive",
                "description": "Generated by performance load tester",
                "parser_config": {"layout_recognize": "DeepDOC"},
            },
        )
        parsing_options_raw = parse_json_object(form.get("parsing_options"), "Parsing config", {"poll_interval_sec": 2, "chunk_sample_size": 3, "collect_pipeline_logs": True})
        retrieval_options_raw = parse_json_object(form.get("retrieval_options"), "Retrieval config", {"concurrency": 4, "top_k": 8, "similarity_threshold": 0.2, "vector_similarity_weight": 0.7, "highlight": True})
        chat_options_raw = parse_json_object(
            form.get("chat_options"),
            "Chat config",
            {"concurrency": 4, "create": {"llm": {"model_name": None}}, "completion": {"extra_body": {"reference": True}}},
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400

    enable_parsing = form.get("enable_parsing") == "on"
    enable_retrieval = form.get("enable_retrieval") == "on"
    enable_chat = form.get("enable_chat") == "on"

    if (enable_retrieval or enable_chat) and not enable_parsing:
        return jsonify({"error": "Retrieval and chat require parsing to be enabled in this upload-based workflow."}), 400

    parsing_request_options = deepcopy(parsing_options_raw)
    parsing_poll_interval = max(1, safe_int(parsing_request_options.pop("poll_interval_sec"), 2))
    chunk_sample_size = max(3, safe_int(parsing_request_options.pop("chunk_sample_size"), 3))
    collect_pipeline_logs = parsing_request_options.pop("collect_pipeline_logs", True)

    retrieval_options = deepcopy(retrieval_options_raw)
    retrieval_concurrency = max(1, safe_int(retrieval_options.pop("concurrency"), 4))
    run_count = min(50, max(1, safe_int(form.get("parallel_runs"), 1)))
    parallel_runs = run_count

    chat_options = deepcopy(chat_options_raw)
    chat_concurrency = max(1, safe_int(chat_options.pop("concurrency"), 4))
    chat_create_options = chat_options.pop("create", {})
    chat_completion_options = chat_options.pop("completion", {})
    if not isinstance(chat_create_options, dict) or not isinstance(chat_completion_options, dict):
        return jsonify({"error": "Chat config must contain object values for 'create' and 'completion'."}), 400
    verify_ssl = form.get("verify_ssl") == "on"
    enable_llm_summary = form.get("enable_llm_summary") == "on"

    batch_id = str(uuid.uuid4())
    batch_root = UPLOADS_DIR / f"batch-{batch_id}"
    source_dir = batch_root / "source"
    source_dir.mkdir(parents=True, exist_ok=True)
    source_paths: list[Path] = []
    filenames: list[str] = []
    run_ids: list[str] = []

    try:
        for file_storage in uploaded_files:
            filename = sanitize_filename(file_storage.filename or f"upload-{uuid.uuid4().hex[:8]}")
            target = source_dir / filename
            await file_storage.save(target)
            source_paths.append(target)
            filenames.append(filename)

        tls_config = {
            "verify": verify_ssl,
        }

        common_config = {
            "base_url": base_url,
            "api_key": api_key,
            "api_key_hint": f"...{api_key[-4:]}" if len(api_key) >= 4 else "***",
            "dataset_options": dataset_options,
            "prompts_per_document": max(1, safe_int(prompt_options.get("prompts_per_document"), 3)),
            "shared_prompts": max(0, safe_int(prompt_options.get("shared_prompts"), 2)),
            "http_timeout_sec": max(30, safe_int(form.get("http_timeout_sec"), 180)),
            "dataset_prefix": form.get("dataset_prefix", "perf-dataset").strip() or "perf-dataset",
            "chat_prefix": form.get("chat_prefix", "perf-chat").strip() or "perf-chat",
            "cleanup_remote": form.get("cleanup_remote") == "on",
            "tls": tls_config,
            "batch_root": str(batch_root),
            "analysis": {"enabled": enable_llm_summary},
            "stages": {
                "parsing": {
                    "enabled": enable_parsing,
                    "poll_interval_sec": parsing_poll_interval,
                    "chunk_sample_size": chunk_sample_size,
                    "collect_pipeline_logs": bool(collect_pipeline_logs),
                    "request_options": parsing_request_options,
                    "raw": parsing_options_raw,
                },
                "retrieval": {
                    "enabled": enable_retrieval,
                    "concurrency": retrieval_concurrency,
                    "request_options": retrieval_options,
                    "raw": retrieval_options_raw,
                },
                "chat": {
                    "enabled": enable_chat,
                    "concurrency": chat_concurrency,
                    "create_options": chat_create_options,
                    "completion_options": chat_completion_options,
                    "raw": chat_options_raw,
                },
            },
        }

        for run_index in range(1, run_count + 1):
            run_id = str(uuid.uuid4())
            upload_dir = batch_root / run_id
            upload_dir.mkdir(parents=True, exist_ok=True)
            local_paths = []
            for source_path in source_paths:
                target = upload_dir / source_path.name
                shutil.copy2(source_path, target)
                local_paths.append(str(target))
            config = deepcopy(common_config)
            config["local_files"] = local_paths
            config["upload_dir"] = str(upload_dir)
            config["batch"] = {"batch_id": batch_id, "run_index": run_index, "run_count": run_count, "parallel_count": parallel_runs}
            queue_run(run_id, config, filenames)
            run_ids.append(run_id)

        return jsonify({"run_id": run_ids[0], "run_ids": run_ids})
    finally:
        shutil.rmtree(source_dir, ignore_errors=True)
        if not run_ids:
            shutil.rmtree(batch_root, ignore_errors=True)


load_existing_runs()


if __name__ == "__main__":
    port = int(os.getenv("PERFORMANCE_APP_PORT", "8787"))
    app.run(host="0.0.0.0", port=port, debug=False)

